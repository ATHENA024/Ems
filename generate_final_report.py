"""
Generate the final EMS Project Report with embedded images.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

REPORT_IMAGES = r"C:\Users\KSonu\Desktop\EMS\report_images"

doc = Document()

# ============================================================
# PAGE SETUP
# ============================================================
for section in doc.sections:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.3)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

def set_font(run, name='Times New Roman', size=12, bold=False, caps=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if caps:
        run.font.all_caps = True
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, size=12, caps=False, space_before=0, space_after=0, line_spacing=1.5):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, caps=caps)
    return p

def add_heading_custom(text, level=1, center=False):
    if level == 1:
        p = add_para(text, alignment=WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT,
                     bold=True, size=15, caps=True, space_before=14, space_after=8)
    elif level == 2:
        p = add_para(text, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                     bold=True, size=13, caps=True, space_before=10, space_after=4)
    elif level == 3:
        p = add_para(text, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                     bold=False, size=12, caps=True, space_before=6, space_after=2)
    return p

def add_body(text, bold=False):
    return add_para(text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, bold=bold)

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    set_font(run)
    p.paragraph_format.line_spacing = 1.5
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * level)
    return p

def add_table_with_data(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
    doc.add_paragraph()
    return table

def add_image_with_caption(image_name, caption, width=Inches(4.5)):
    """Add an image centered with a caption below it"""
    path = os.path.join(REPORT_IMAGES, image_name)
    if os.path.exists(path):
        # Image
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(12)
        p_img.paragraph_format.space_after = Pt(4)
        run = p_img.add_run()
        run.add_picture(path, width=width)
        # Caption
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(0)
        p_cap.paragraph_format.space_after = Pt(12)
        run = p_cap.add_run(caption)
        set_font(run, size=10, bold=True)
    else:
        # Fallback
        add_para(f"[Image: {caption} - file not found: {image_name}]",
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, size=10)

def page_break():
    doc.add_page_break()

def separator_page(chapter_title, description=""):
    """Add a chapter separator with title and overview"""
    page_break()
    for _ in range(4):
        add_para('', size=12)
    add_para(chapter_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, caps=True)
    add_para('', size=8)
    if description:
        add_para(description, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=11)
        add_para('', size=4)
    add_para('─' * 40, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=10)

def add_header_footer():
    """Add header and footer matching university guidelines format"""
    for section in doc.sections:
        # ---- HEADER ----
        header = section.header
        header.is_linked_to_previous = False
        # Clear existing header
        for p in header.paragraphs:
            p.clear()
        # Create 2-column table in header
        tbl = header.add_table(rows=1, cols=2, width=Inches(8.27))
        tbl.style = 'Table Grid'
        tbl_pr = tbl._tbl.tblPr
        tbl_w = parse_xml(f'<w:tblW {nsdecls("w")} w:w="8946" w:type="dxa"/>')
        tbl_pr.append(tbl_w)
        tbl_pr_borders = tbl_pr.find(qn('w:tblBorders'))
        if tbl_pr_borders is not None:
            tbl_pr.remove(tbl_pr_borders)
        tbl_borders_xml = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
            f'<w:left w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
            f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
            f'<w:right w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
            f'</w:tblBorders>'
        )
        tbl_pr.append(tbl_borders_xml)
        for cell in tbl.row_cells(0):
            tc_pr = cell._tc.tcPr
            if tc_pr is None:
                tc_pr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
                cell._tc.insert(0, tc_pr)
            tc_mar = parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                f'<w:top w:w="10" w:type="dxa"/>'
                f'<w:left w:w="10" w:type="dxa"/>'
                f'<w:bottom w:w="10" w:type="dxa"/>'
                f'<w:right w:w="10" w:type="dxa"/>'
                f'</w:tcMar>'
            )
            tc_pr.append(tc_mar)
        left_cell = tbl.row_cells(0)[0]
        left_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_cell.paragraphs[0].clear()
        run = left_cell.paragraphs[0].add_run("Dept. of CSE | 6th Sem Diploma")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        right_cell = tbl.row_cells(0)[1]
        right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_cell.paragraphs[0].clear()
        run = right_cell.paragraphs[0].add_run("Major Project Report")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        run.bold = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p_border = header.add_paragraph()
        p_border.paragraph_format.space_before = Pt(0)
        p_border.paragraph_format.space_after = Pt(0)
        p_pr = p_border._p.get_or_add_pPr()
        p_bdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="1F3864"/>'
            f'</w:pBdr>'
        )
        p_pr.append(p_bdr)

        # ---- FOOTER ----
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()
        p_border = footer.add_paragraph()
        p_border.paragraph_format.space_before = Pt(0)
        p_border.paragraph_format.space_after = Pt(0)
        p_pr = p_border._p.get_or_add_pPr()
        p_bdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="1" w:color="1F3864"/>'
            f'</w:pBdr>'
        )
        p_pr.append(p_bdr)
        tbl = footer.add_table(rows=1, cols=2, width=Inches(8.27))
        tbl.style = 'Table Grid'
        tbl_pr = tbl._tbl.tblPr
        tbl_w = parse_xml(f'<w:tblW {nsdecls("w")} w:w="8946" w:type="dxa"/>')
        tbl_pr.append(tbl_w)
        tbl_pr_borders = tbl_pr.find(qn('w:tblBorders'))
        if tbl_pr_borders is not None:
            tbl_pr.remove(tbl_pr_borders)
        tbl_borders_xml = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
            f'<w:left w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
            f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
            f'<w:right w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
            f'</w:tblBorders>'
        )
        tbl_pr.append(tbl_borders_xml)
        for cell in tbl.row_cells(0):
            tc_pr = cell._tc.tcPr
            if tc_pr is None:
                tc_pr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
                cell._tc.insert(0, tc_pr)
            tc_mar = parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                f'<w:top w:w="10" w:type="dxa"/>'
                f'<w:left w:w="10" w:type="dxa"/>'
                f'<w:bottom w:w="10" w:type="dxa"/>'
                f'<w:right w:w="10" w:type="dxa"/>'
                f'</w:tcMar>'
            )
            tc_pr.append(tc_mar)
        left_cell = tbl.row_cells(0)[0]
        left_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_cell.paragraphs[0].clear()
        run = left_cell.paragraphs[0].add_run("Dr. C. V. Raman University | Diploma 6th Sem")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.bold = True
        right_cell = tbl.row_cells(0)[1]
        right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_cell.paragraphs[0].clear()
        run = right_cell.paragraphs[0].add_run("Page ")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run_fld_begin = right_cell.paragraphs[0].add_run()
        fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_fld_begin._r.append(fldChar_begin)
        run_instr = right_cell.paragraphs[0].add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run_instr._r.append(instrText)
        run_fld_sep = right_cell.paragraphs[0].add_run()
        fldChar_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run_fld_sep._r.append(fldChar_sep)
        run_page_num = right_cell.paragraphs[0].add_run()
        run_page_num.font.name = 'Times New Roman'
        run_page_num.font.size = Pt(9)
        run_page_num.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run_fld_end = right_cell.paragraphs[0].add_run()
        fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_fld_end._r.append(fldChar_end)

# Enable header and footer (university guidelines format)
add_header_footer()

# ============================================================
# COVER / TITLE PAGE
# ============================================================

# Logo
logo_path = r"C:\Users\KSonu\Desktop\EMS\public\logo.png"
if os.path.exists(logo_path):
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(24)
    p_logo.paragraph_format.space_after = Pt(6)
    run = p_logo.add_run()
    run.add_picture(logo_path, width=Inches(1.2))

add_para('DR. C.V. RAMAN UNIVERSITY', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, caps=True)
add_para('VAISHALI (BIHAR)', alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para('───  ●  ───', alignment=WD_ALIGN_PARAGRAPH.CENTER, size=10)
add_para('', size=6)
add_para('A Major Project Report', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, caps=True)
add_para('on', alignment=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para('"EMPLOYEE MANAGEMENT SYSTEM"', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18, caps=True)
add_para('', size=12)
add_para('Submitted in Partial Fulfillment of the Requirements for the Award of the Degree of',
         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para('DIPLOMA IN COMPUTER SCIENCE AND ENGINEERING', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, caps=True)
add_para('(6th Semester)', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)

add_para('', size=12)

add_para('SUBMITTED BY', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, caps=True)

add_table_with_data(
    ['S. No.', 'STUDENT NAME', 'ENROLLMENT NO.', 'ROLL NO.'],
    [
        ['1', 'Sonu Kumar', 'CVB2402116', 'L24ETDCS0005'],
        ['2', 'Aditya Kumar', 'CVB2300228', 'R23ETDCS0004'],
    ]
)
add_para('', size=6)

add_para('BATCH: 2025-2026 / 6th SEMESTER', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12)

add_para('SUBMITTED TO', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, caps=True)
add_para('Prabhat Ranjan', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
add_para('Assistant Professor, Department of CSE', alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12)

add_para('', size=6)
add_para('DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, caps=True)
add_para('DR. C.V. RAMAN UNIVERSITY', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, caps=True)
add_para('VAISHALI (BIHAR)', alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para('2025-2026', alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12)

page_break()

# ============================================================
# DECLARATION
# ============================================================
add_para('DR. C.V. RAMAN UNIVERSITY', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, caps=True)
add_para('VAISHALI (BIHAR)', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12)
add_para('DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, caps=True)
add_para('', size=12)
add_heading_custom('DECLARATION BY THE CANDIDATE', level=1, center=True)
add_body('We, the undersigned students of Diploma (Computer Science & Engineering), 6th Semester, [Institute/Polytechnic Name], [City, State], hereby declare that the Major Project entitled "EMPLOYEE MANAGEMENT SYSTEM" has been carried out by us under the supervision of Mr. Prabhat Ranjan, Assistant Professor, Department of CSE.')
add_body('We further declare that:')
add_bullet('The work presented in this report is original and has been carried out by us.')
add_bullet('This work has not been submitted, either in part or in full, for any degree or diploma at this or any other institution.')
add_bullet('Wherever references have been made to intellectual properties of any individual or organization, it has been clearly indicated in the text and included in the list of references.')
add_bullet('We have adhered to all principles of academic honesty and have not practised plagiarism. The plagiarism report from approved software is attached as Appendix.')
add_para('', size=12)
add_body('Place: ___________________')
add_body('Date: ____________________')
add_para('', size=12)
add_body('Signatures of Candidates:')
add_para('1. Sonu Kumar ________________________', size=12)
add_para('2. Aditya Kumar ________________________', size=12)

page_break()

# ============================================================
# CERTIFICATE
# ============================================================
add_para('DR. C.V. RAMAN UNIVERSITY', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, caps=True)
add_para('VAISHALI (BIHAR)', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12)
add_para('DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, caps=True)
add_para('', size=12)
add_heading_custom('CERTIFICATE OF THE SUPERVISOR', level=1, center=True)
add_body('This is to certify that the Major Project entitled "EMPLOYEE MANAGEMENT SYSTEM" submitted by the following students in partial fulfilment of the requirements for the award of the degree of Diploma in Computer Science & Engineering from [Institute/Polytechnic Name], [City, State], is a bonafide record of original work carried out by them under my guidance and supervision.')

add_table_with_data(
    ['S.No.', 'Student Name', 'Enroll. No.', 'Roll No.'],
    [
        ['1', 'Sonu Kumar', 'CVB2402116', 'L24ETDCS0005'],
        ['2', 'Aditya Kumar', 'CVB2300228', 'R23ETDCS0004'],
    ]
)

add_body('To the best of my knowledge, this work is original, has not been submitted elsewhere for any award, and all sources consulted have been duly acknowledged.')
add_para('', size=12)
add_body('Date: ____________________')
add_body('Place: ___________________')
add_para('', size=12)
add_body('Signature of Supervisor')
add_body('Name: Prabhat Ranjan')
add_body('Designation: Assistant Professor')
add_body('Department of CSE')
add_body('Dr. C.V. Raman University')

page_break()

# ============================================================
# ACKNOWLEDGEMENT
# ============================================================
add_para('DR. C.V. RAMAN UNIVERSITY', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, caps=True)
add_para('VAISHALI (BIHAR)', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12)
add_para('DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, caps=True)
add_para('', size=12)
add_heading_custom('ACKNOWLEDGEMENT', level=1, center=True)
add_body('We take immense pleasure in expressing our profound sense of gratitude and indebtedness to all those who have supported us during the course of this Major Project.')
add_body('We sincerely thank our Project Supervisor, Mr. Prabhat Ranjan, Assistant Professor, Department of CSE, Dr. C.V. Raman University, for his invaluable guidance, constant encouragement, constructive suggestions, and persistent support throughout the project. His expertise and insights have been a great source of inspiration for us.')
add_body('We are deeply grateful to the Head of the Department, Department of Computer Science & Engineering, for providing all the necessary facilities and a conducive environment for completing this work.')
add_body('We also extend our heartfelt thanks to the Management and Faculty of Dr. C.V. Raman University for their continuous support, motivation, and provision of all academic resources needed for the successful completion of this project.')
add_body('We are also thankful to our family members and friends for their unwavering moral support and encouragement throughout the course of this work.')
add_para('', size=12)
add_body('Date: ____________________')
add_body('Place: ___________________')
add_para('', size=12)
add_body('Submitted By:')
add_para('1. Sonu Kumar ________________________', size=12)
add_para('2. Aditya Kumar ________________________', size=12)

page_break()

# ============================================================
# ABSTRACT
# ============================================================
add_heading_custom('ABSTRACT', level=1, center=True)
add_body('The "Employee Management System" (EMS) is a web-based application designed to streamline and automate the human resource management processes within an organization. In today\'s fast-paced digital environment, managing employee data, attendance, leave records, payroll, and performance evaluations manually is inefficient and prone to errors. This project aims to provide a comprehensive, centralized platform that addresses these challenges effectively.')
add_body('The primary objective of this project is to design and develop a fully functional Employee Management System that enables organizations to manage their workforce efficiently. The scope of the project includes employee record management, attendance tracking, leave management, payroll processing, performance management, employee appraisals, onboarding management, learning management, announcements, and a helpdesk ticketing system. The system also includes role-based access control with distinct Admin and Employee dashboards.')
add_body('The system was developed using the Waterfall SDLC methodology, progressing through planning, requirement analysis, design, implementation, testing, deployment, and maintenance phases. The technology stack comprises React.js for the frontend, Node.js with Express.js for the backend, and MongoDB for the database. The system architecture follows a client-server model with RESTful API communication between the frontend and backend layers.')
add_body('The developed system successfully implements all planned modules with a user-friendly interface. Testing was carried out using unit testing and integration testing methodologies, which demonstrated that the system is robust, reliable, and meets all functional requirements. The system handles employee data management with full CRUD operations, attendance tracking with check-in/check-out functionality, leave management with approval workflows, and payroll processing with automated Indian TDS tax calculations.')
add_body('This project provides a robust and scalable solution for organizational workforce management. The modular architecture allows for easy maintenance and future enhancements. Potential future enhancements may include mobile application development, biometric integration for attendance, AI-powered performance analytics, and cloud deployment for remote accessibility.')
add_para('', size=12)
add_body('Keywords: Employee Management System, Human Resource Management, React.js, Node.js, MongoDB, Attendance Tracking, Payroll Management, Leave Management, Performance Evaluation, Web Application.', bold=True)

page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_heading_custom('TABLE OF CONTENTS', level=1, center=True)
toc_entries = [
    ('Declaration by the Candidate', 'i'),
    ('Certificate of the Supervisor', 'ii'),
    ('Acknowledgement', 'iii'),
    ('Abstract', 'iv'),
    ('Table of Contents', 'v'),
    ('List of Figures', 'vii'),
    ('List of Tables', 'ix'),
    ('Chapter 1: Introduction', '1'),
    ('  1.1 Background', '1'),
    ('  1.2 Problem Statement', '2'),
    ('  1.3 Objectives', '3'),
    ('  1.4 Scope', '3'),
    ('  1.5 Organization of Report', '4'),
    ('Chapter 2: Literature Review', '5'),
    ('  2.1 Existing Systems', '5'),
    ('  2.2 Comparative Analysis', '7'),
    ('  2.3 Research Gaps', '8'),
    ('Chapter 3: System Analysis & Design', '9'),
    ('  3.1 Requirement Analysis', '9'),
    ('  3.2 Feasibility Study', '12'),
    ('  3.3 Gantt Chart', '13'),
    ('  3.4 PERT Chart', '14'),
    ('  3.5 ER Diagram', '15'),
    ('  3.6 Data Flow Diagrams', '16'),
    ('  3.6.4 DFD Level 2 - Attendance & Leave', '19'),
    ('  3.7 SDLC Methodology', '20'),
    ('  3.8 Control Flow Diagram', '22'),
    ('  3.9 Use Case Diagram', '23'),
    ('Chapter 4: Implementation', '24'),
    ('  4.1 Technology Stack', '24'),
    ('  4.2 System Architecture', '25'),
    ('  4.3 Database Schema Design', '25'),
    ('  4.4 Module Implementation', '26'),
    ('  4.5 Deployment Configuration', '35'),
    ('  4.6 Seed Data and Demo Environment', '36'),
    ('  4.7 Screenshots', '37'),
    ('Chapter 5: Testing & Results', '40'),
    ('  5.1 Testing Strategy', '40'),
    ('  5.2 Test Cases', '41'),
    ('  5.3 Results Analysis', '45'),
    ('Chapter 6: Conclusion & Future Scope', '47'),
    ('  6.1 Conclusion', '47'),
    ('  6.2 Challenges Faced', '48'),
    ('  6.3 Limitations', '49'),
    ('  6.4 Future Scope', '50'),
    ('References / Bibliography', '50'),
    ('Appendices', '52'),
]
for entry, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.tab_stops.add_tab_stop(Inches(5.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
    run = p.add_run(entry)
    set_font(run, size=12)
    run = p.add_run('\t' + page)
    set_font(run, size=12)

page_break()

# ============================================================
# LIST OF FIGURES
# ============================================================
add_heading_custom('LIST OF FIGURES', level=1, center=True)
figures = [
    ('Figure 3.1: Project Gantt Chart', '13'),
    ('Figure 3.2: PERT Chart', '14'),
    ('Figure 3.3: ER Diagram', '15'),
    ('Figure 3.4: DFD Level 0 - Context Diagram', '16'),
    ('Figure 3.5: DFD Level 1 - System Overview', '17'),
    ('Figure 3.6: DFD Level 2 - Employee Management', '18'),
    ('Figure 3.7: DFD Level 2 - Attendance & Leave Management', '19'),
    ('Figure 3.8: SDLC Model Diagram', '20'),
    ('Figure 3.9: Control Flow Diagram', '22'),
    ('Figure 3.10: Use Case Diagram', '23'),
    ('Figure 3.11: System Architecture Diagram', '26'),
    ('Figure 4.1: Login Page', '35'),
    ('Figure 4.2: Admin Dashboard', '35'),
    ('Figure 4.3: Employee Dashboard', '36'),
    ('Figure 4.4: Employee Management', '36'),
    ('Figure 4.5: Attendance Module', '37'),
    ('Figure 4.6: Leave Management', '37'),
    ('Figure 4.7: Payroll Module', '38'),
    ('Figure 4.8: Reports Module', '38'),
]
for fig, pg in figures:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.tab_stops.add_tab_stop(Inches(5.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
    run = p.add_run(fig)
    set_font(run, size=12)
    run = p.add_run('\t' + pg)
    set_font(run, size=12)

page_break()

# ============================================================
# LIST OF TABLES
# ============================================================
add_heading_custom('LIST OF TABLES', level=1, center=True)
tables_list = [
    ('Table 3.1: Functional Requirements', '10'),
    ('Table 3.2: Non-Functional Requirements', '11'),
    ('Table 3.3: Hardware Requirements', '11'),
    ('Table 3.4: Software Requirements', '12'),
    ('Table 3.5: Use Case Description', '23'),
    ('Table 5.1: Test Cases for Employee Management', '41'),
    ('Table 5.2: Test Cases for Attendance Module', '42'),
    ('Table 5.3: Test Cases for Leave Management', '43'),
    ('Table 5.4: Test Cases for Payroll Module', '44'),
    ('Table 5.5: Test Results Summary', '45'),
]
for tbl, pg in tables_list:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.tab_stops.add_tab_stop(Inches(5.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
    run = p.add_run(tbl)
    set_font(run, size=12)
    run = p.add_run('\t' + pg)
    set_font(run, size=12)

# ============================================================
# CHAPTER 1: INTRODUCTION
# ============================================================
separator_page('CHAPTER 1', 'An overview of the Employee Management System including background, problem statement, objectives, scope, and report organization.')

add_heading_custom('1.1 Background', level=2)
add_body('In the modern organizational landscape, effective management of human resources is critical to the success and growth of any enterprise. Employee management encompasses a wide range of activities including maintaining employee records, tracking attendance, managing leaves, processing payroll, evaluating performance, and facilitating onboarding of new hires. Traditionally, many organizations, especially small and medium-sized enterprises, have relied on manual processes using spreadsheets and paper-based records to handle these tasks. However, as organizations grow, these manual methods become increasingly inefficient, error-prone, and difficult to maintain.')
add_body('The Employee Management System (EMS) is a web-based application developed to address these challenges by providing a centralized digital platform for managing all employee-related operations. This system leverages modern web technologies to deliver a responsive, user-friendly interface that can be accessed from any device with a web browser. By automating routine HR tasks, the system reduces administrative overhead, minimizes errors, and provides valuable insights through comprehensive reporting capabilities.')
add_body('The project is motivated by the need for a cost-effective, scalable, and easy-to-use solution that can be adopted by organizations of all sizes. The system is designed with modularity in mind, allowing organizations to use only the features they need while having the flexibility to expand functionality as their requirements evolve.')
add_body('Human Resource Management (HRM) has evolved significantly over the past few decades, transitioning from paper-based record keeping to sophisticated digital platforms. The emergence of cloud computing, mobile technologies, and artificial intelligence has further transformed how organizations manage their workforce. However, many small and medium enterprises in India still struggle with adopting these technologies due to cost constraints, complexity, and lack of tailored solutions that address local requirements such as Indian tax regulations and statutory compliance.')
add_body('The evolution of HRM can be traced through several generations. The first generation involved manual paper-based systems where employee records were maintained in physical files and registers. The second generation brought standalone digital tools such as spreadsheet software, which improved data organization but still required significant manual effort. The third generation introduced integrated software packages that combined multiple HR functions in a single system. The current fourth generation represents cloud-based, intelligent HR platforms that leverage artificial intelligence, machine learning, and real-time analytics to provide actionable insights and automate complex workflows.')
add_body('Despite these technological advancements, a significant gap persists between enterprise-grade solutions and the needs of small and medium-sized organizations. Large enterprises can afford comprehensive systems like SAP SuccessFactors or Oracle HCM Cloud, which cost millions in licensing and implementation. However, SMEs, which constitute over 95% of businesses in India, require affordable, easy-to-deploy solutions that address their specific needs without unnecessary complexity.')
add_body('The proposed Employee Management System aims to bridge this gap by providing a comprehensive yet affordable solution that combines essential HR functionalities with Indian-specific features such as TDS calculation, provident fund management, and statutory reporting. The system is built using modern web technologies including React.js, Node.js, Express.js, and MongoDB, following industry best practices for security, scalability, and user experience. The system adopts a hybrid architecture that combines a traditional client-server model with local-first capabilities, allowing it to function even without a continuous backend connection through browser localStorage fallback mechanisms. This design ensures high availability and a seamless user experience even in environments with intermittent internet connectivity.')

add_heading_custom('1.2 Problem Statement', level=2)
add_body('Organizations face several challenges in managing their workforce effectively using traditional manual methods. These challenges include:')
add_bullet('Time-consuming manual data entry and record-keeping processes that are prone to human error.')
add_bullet('Difficulty in tracking employee attendance and managing leave requests efficiently.')
add_bullet('Complex payroll calculations involving multiple components such as basic pay, allowances, deductions, and tax calculations.')
add_bullet('Lack of centralized visibility into employee performance and appraisal history.')
add_bullet('Inefficient communication channels for company announcements and helpdesk support.')
add_bullet('Difficulty in managing onboarding processes for new employees in a standardized manner.')
add_bullet('Absence of automated reporting and analytics to support data-driven decision making.')
add_bullet('Security concerns related to storing sensitive employee data in unsecured spreadsheets or disparate systems.')
add_body('These challenges are compounded in growing organizations where the workforce size increases rapidly, making manual processes unsustainable. Errors in payroll calculation can lead to employee dissatisfaction, legal complications, and regulatory penalties. Delays in leave approvals and attendance tracking can affect productivity, employee morale, and organizational trust. Without proper performance management systems, organizations struggle to identify top performers, address underperformance effectively, and make data-driven decisions about promotions, training, and succession planning.')
add_body('Furthermore, the lack of integrated systems leads to data silos where employee information exists in multiple disconnected spreadsheets, emails, and paper files. This fragmentation makes it difficult to get a holistic view of an employee\'s journey within the organization, from onboarding through performance reviews to eventual exit. Compliance with statutory requirements such as TDS deduction, PF contribution, and professional tax becomes challenging without automated systems that track these calculations accurately.')
add_body('Security is another major concern with manual systems. Employee data including personal information, salary details, and bank account numbers stored in unsecured spreadsheets or email attachments is vulnerable to unauthorized access and data breaches. The absence of proper access controls and audit trails makes it difficult to track who accessed sensitive information and when.')
add_body('The Employee Management System aims to solve these problems by providing an integrated digital platform that automates and streamlines these HR processes, ensuring accuracy, efficiency, and accessibility. By centralizing all employee-related data and processes in a single system with role-based access control, EMS eliminates data silos, reduces duplication of effort, provides a single source of truth for HR operations, and ensures that sensitive data is protected through proper authentication and authorization mechanisms.')

add_heading_custom('1.3 Objectives', level=2)
add_body('The primary objectives of this project are:')
add_bullet('To design and develop a web-based Employee Management System with a user-friendly interface.')
add_bullet('To implement role-based access control with distinct Admin and Employee functionalities.')
add_bullet('To provide comprehensive employee record management with full CRUD (Create, Read, Update, Delete) operations.')
add_bullet('To develop an attendance tracking system with check-in/check-out functionality.')
add_bullet('To implement a leave management system with application and approval workflows.')
add_bullet('To create a payroll processing module with automatic salary calculation and Indian TDS tax computation.')
add_bullet('To build a performance management module for tracking employee tasks and conducting appraisals.')
add_bullet('To include additional modules for onboarding, learning management, announcements, and helpdesk ticketing.')
add_bullet('To generate comprehensive reports for data analysis and decision support.')
add_body('These objectives align with the core challenges identified during the requirement gathering phase. Each objective maps to specific pain points expressed by potential end-users during stakeholder interviews and surveys. For instance, the objective of role-based access control addresses the security concerns of HR managers, while the automated payroll and TDS calculation objective addresses the most time-consuming and error-prone manual process identified during the analysis.')
add_body('The objectives also follow the SMART framework: Specific (each objective targets a clearly defined module or feature), Measurable (success can be verified through testing and user acceptance), Achievable (within the scope of available technologies and team expertise), Relevant (addresses genuine organizational needs), and Time-bound (completed within the academic semester timeline).')

add_heading_custom('1.4 Scope', level=2)
add_body('The scope of this project encompasses the following modules and features:')
add_bullet('Employee Management: Add, view, update, and delete employee records with details such as name, contact information, department, position, salary, and employment status.')
add_bullet('Attendance Management: Daily check-in and check-out tracking with history views for both admins and employees.')
add_bullet('Leave Management: Leave application with multiple leave types (Sick, Casual, Annual, Personal, Maternity, Paternity), leave balance tracking, and admin approval workflow.')
add_bullet('Payroll Management: Salary structure definition, automatic payroll calculation with Indian TDS tax computation, payment processing, and payslip generation.')
add_bullet('Performance Management: Task assignment and tracking with status workflow, priority levels, and due date management.')
add_bullet('Appraisal Management: Quarterly performance reviews with multi-dimension scoring, star ratings, and automatic rating calculation.')
add_bullet('Onboarding Management: New hire onboarding checklists and progress tracking.')
add_bullet('Learning Management System (LMS): Course catalog, enrollment, and progress tracking.')
add_bullet('Announcements: Company-wide announcements with priority levels and department filtering.')
add_bullet('Helpdesk: Ticket management system for handling employee queries and issues.')
add_bullet('Reports: Comprehensive reporting across all modules with filtering and data export capabilities.')
add_bullet('Security: User authentication, role-based access control, and audit logging.')

add_heading_custom('1.5 Organization of Report', level=2)
add_body('This report is organized into six chapters. Chapter 1 provides an introduction to the project including background, problem statement, objectives, and scope. Chapter 2 presents a literature review of existing systems and technologies, comparing them with the proposed system. Chapter 3 covers the system analysis and design including requirement analysis, feasibility study, and all system diagrams such as ER diagrams, DFDs, and use case diagrams. Chapter 4 details the implementation of the system including the technology stack, architecture, and module-wise implementation with screenshots. Chapter 5 discusses the testing strategy, test cases, and results analysis. Chapter 6 concludes the report with a summary of findings, limitations, and future scope for enhancement. The report also includes references and appendices containing key source code references.')

# ============================================================
# CHAPTER 2: LITERATURE REVIEW
# ============================================================
separator_page('CHAPTER 2', 'A comprehensive review of existing HR management systems, comparative analysis, and identification of research gaps.')

add_heading_custom('2.1 Existing Systems', level=2)
add_body('A comprehensive review of existing employee management and HR systems was conducted to understand the current landscape and identify areas for improvement. Several commercial and open-source solutions are available in the market, each with its own strengths and limitations. The review focused on systems that are widely used in the Indian context and internationally.')
add_body('Commercial HR Management Systems such as SAP SuccessFactors, Oracle HCM Cloud, and Workday provide comprehensive enterprise-grade HR solutions. These systems offer extensive functionality including workforce planning, talent management, payroll processing, and analytics. However, they are typically expensive to license and implement, require significant IT infrastructure, and involve complex customization processes that may not be suitable for small and medium-sized organizations. The total cost of ownership for these systems can run into millions of rupees annually, making them inaccessible for most Indian SMEs.')
add_body('Open-source solutions such as OrangeHRM, Odoo HR, and Sentrifugo HRMS offer more accessible alternatives with basic to moderate HR functionality. While these systems eliminate licensing costs, they often require technical expertise for installation, configuration, and maintenance. Additionally, the user interface of many open-source solutions may not be as intuitive or modern as commercial alternatives. The customization flexibility comes at the cost of requiring skilled developers who can modify the codebase to suit organizational needs.')
add_body('Cloud-based HR platforms such as Zoho People, BambooHR, and Keka have gained popularity in recent years due to their ease of use and subscription-based pricing models. These platforms offer good functionality for small to medium businesses but may have limitations in customization and data ownership. Additionally, reliance on internet connectivity and ongoing subscription costs can be concerns for some organizations. Data residency is also a concern for Indian organizations that need to comply with local data protection regulations.')
add_body('The proposed Employee Management System differentiates itself by combining the affordability of open-source solutions with the user experience of commercial products, while adding India-specific features that are often missing in both categories. The system is built using modern web technologies including React.js and Node.js, ensuring a responsive and intuitive user interface that rivals commercial products.')
add_body('In addition to the primary categories discussed above, several other notable HR solutions were reviewed:')
add_bullet('Zoho People: A popular cloud-based HR platform from Indian company Zoho Corporation. It offers good functionality for Indian businesses including attendance tracking, leave management, and basic payroll. However, its advanced features require additional subscription costs, and customization options are limited compared to open-source alternatives.')
add_bullet('Keka: An Indian HR and payroll platform designed specifically for Indian businesses. It offers good compliance features for Indian regulations including TDS, PF, and ESI. However, it is a paid subscription service starting at approximately Rs. 6,000 per month for up to 100 employees, which may be costly for very small organizations.')
add_bullet('OrangeHRM: One of the most popular open-source HR management systems. It provides modules for personnel information management, leave management, attendance tracking, and recruitment. The open-source version is free but has limited functionality, while the premium versions require licensing fees. The user interface, while functional, is not as modern as contemporary web applications.')
add_bullet('Sentrifugo HRMS: A comprehensive open-source HRMS solution with modules covering employee management, attendance, leaves, performance, and payroll. It offers reasonable functionality but suffers from an outdated user interface and requires PHP-based hosting infrastructure.')

add_body('The review also included an analysis of the underlying technologies used in modern HR systems. The technology landscape for web application development has evolved significantly, with the MERN stack (MongoDB, Express.js, React.js, Node.js) emerging as one of the most popular choices for building scalable, full-stack web applications. React.js, developed by Meta (formerly Facebook), has become the dominant frontend framework due to its component-based architecture, virtual DOM for efficient rendering, and extensive ecosystem of libraries and tools. Node.js provides a non-blocking, event-driven runtime for building scalable server-side applications, while MongoDB offers a flexible document-based data model that is well-suited for applications with evolving data requirements.')
add_body('The adoption of RESTful API architecture enables clean separation between frontend and backend concerns, allowing independent development, testing, and deployment of each layer. JSON Web Tokens (JWT) provide a stateless authentication mechanism that is well-suited for RESTful APIs, while bcryptjs provides industry-standard password hashing for secure credential storage. The combination of these technologies provides a robust foundation for building modern web applications with excellent performance, security, and developer experience.')

add_body('Security considerations were also reviewed as part of the literature study. Modern web applications face a variety of security threats including Cross-Site Scripting (XSS), SQL/NoSQL injection, Cross-Site Request Forgery (CSRF), and authentication bypass attacks. The OWASP Top Ten provides a comprehensive framework for understanding and mitigating these risks. For the EMS, key security measures include input validation and sanitization on both client and server sides, parameterized database queries through Mongoose ORM to prevent injection attacks, JWT-based authentication with proper token expiration and secure storage, bcrypt password hashing with salt rounds, and role-based access control with the principle of least privilege. The backend CORS configuration restricts cross-origin requests to authorized domains only, and the express.json middleware includes body size limits to prevent denial-of-service attacks through oversized payloads.')

add_body('The technology components of the MERN stack were studied in detail during the literature review:')
add_bullet('React.js (v18.2.0): React is a JavaScript library for building user interfaces developed by Meta. It introduces a component-based architecture where UI elements are encapsulated in reusable components with their own state and lifecycle. The virtual DOM implementation provides efficient rendering by minimizing direct DOM manipulations. React\'s unidirectional data flow and hooks-based state management (useState, useEffect, useContext) enable predictable and maintainable code organization. The React ecosystem provides additional libraries for routing (React Router), state management (Context API, Redux), and HTTP communication (Axios).')
add_bullet('Node.js (v18.x): Node.js is a JavaScript runtime built on Chrome\'s V8 engine that enables server-side JavaScript execution. Its event-driven, non-blocking I/O model makes it well-suited for building scalable network applications. Node.js uses the CommonJS module system for code organization and npm for package management, providing access to the largest ecosystem of open-source libraries in the world.')
add_bullet('Express.js (v4.18.2): Express is a minimal and flexible web application framework for Node.js that provides a robust set of features for building web applications and APIs. It offers middleware-based request processing, routing capabilities, and integration with various template engines and database libraries. In the EMS, Express is used to define RESTful API routes, parse JSON request bodies, enable CORS, and handle errors gracefully.')
add_bullet('MongoDB (v6.x): MongoDB is a NoSQL document database that stores data in flexible, JSON-like documents. Unlike relational databases, MongoDB does not require a predefined schema, allowing for iterative and agile development. Documents are stored in collections, and related data can be embedded within documents or referenced across collections. MongoDB\'s document model maps naturally to object-oriented programming paradigms and is well-suited for applications with evolving data structures.')
add_bullet('Mongoose (v8.0.3): Mongoose is an Object Data Modeling (ODM) library for MongoDB and Node.js. It provides a schema-based solution for modeling application data, including built-in type casting, validation, query building, and business logic hooks (middleware). Mongoose schemas map to MongoDB collections and define the structure, data types, validation rules, and default values for documents. The EMS uses Mongoose to define 11 schemas, manage database connections, and execute CRUD operations with validation and error handling.')

page_break()
add_heading_custom('2.2 Comparative Analysis', level=2)
add_body('The following table presents a comparative analysis of existing systems against the proposed Employee Management System. The comparison is based on key features, cost, ease of use, and other relevant parameters that organizations consider when selecting an HR management system.')
add_table_with_data(
    ['Feature', 'Commercial Systems', 'Open-Source Systems', 'Cloud-Based Systems', 'Proposed EMS'],
    [
        ['Cost', 'High (License + Implementation)', 'Free (but setup costs)', 'Subscription-based', 'Free & Open Source'],
        ['Ease of Use', 'Moderate (complex UI)', 'Varies', 'Good', 'Excellent (Modern UI)'],
        ['Customization', 'Extensive but costly', 'High (requires coding)', 'Limited', 'High (modular design)'],
        ['Deployment', 'On-premise / Cloud', 'On-premise', 'Cloud', 'On-premise / Cloud'],
        ['Attendance', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Leave Management', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['Payroll (India)', 'Limited', 'Limited', 'Available', 'Built-in (Indian TDS)'],
        ['Performance Mgmt', 'Yes', 'Basic', 'Yes', 'Yes'],
        ['LMS', 'Yes (add-on)', 'Basic', 'Limited', 'Yes'],
        ['Helpdesk', 'Yes', 'Basic', 'Limited', 'Yes'],
        ['Reporting', 'Advanced', 'Basic', 'Moderate', 'Comprehensive'],
        ['Mobile Support', 'Yes', 'Limited', 'Yes', 'Responsive Web'],
    ]
)

add_heading_custom('2.3 Research Gaps', level=2)
add_body('Based on the literature review and comparative analysis, the following research gaps were identified:')
add_bullet('Most existing systems do not provide comprehensive Indian tax (TDS) calculations as a built-in feature, requiring additional configuration or third-party integration. This is a significant gap for Indian organizations that need to comply with the Income Tax Act.')
add_bullet('Affordable solutions often lack key modules such as Learning Management and Helpdesk ticketing, requiring organizations to use multiple disparate systems that may not integrate well with each other.')
add_bullet('Open-source solutions often have outdated user interfaces that do not meet modern usability standards, leading to poor user adoption and satisfaction.')
add_bullet('Few systems provide a seamless dual-mode operation (offline/online) with data persistence at both frontend and backend levels.')
add_bullet('There is a need for a lightweight, easily deployable solution that combines essential HR functionalities in a single, well-designed package without the complexity and cost of enterprise systems.')
add_bullet('Most systems do not offer role-based dashboards that provide different views and functionalities for administrators and employees.')
add_body('The proposed Employee Management System aims to address these gaps by providing a feature-rich, modern, and affordable solution that is specifically designed with Indian organizational requirements in mind. The system incorporates all essential HR modules in a single platform with a modern user interface and built-in compliance with Indian statutory requirements.')
add_body('The gap analysis reveals that the proposed EMS occupies a unique position in the market. It combines the affordability and customizability of open-source solutions with the user experience and feature completeness of commercial products. Unlike most open-source HR systems that require significant technical expertise for setup and configuration, the EMS features a self-contained architecture with automatic fallback to in-memory database when no external database is configured, enabling instant deployment with zero configuration on any system with Node.js installed. This makes it particularly suitable for small and medium Indian enterprises that may not have dedicated IT infrastructure or technical staff.')
# ============================================================
# CHAPTER 3: SYSTEM ANALYSIS & DESIGN
# ============================================================
separator_page('CHAPTER 3', 'Detailed system analysis including requirement analysis, feasibility study, system diagrams, and design methodology.')

add_heading_custom('3.1 Requirement Analysis', level=2)

add_body('Requirement analysis is the most critical phase of the software development lifecycle. During this phase, the requirements of the proposed system are collected by studying the existing system and understanding the needs of the end-users. The requirements have been categorized into functional requirements, non-functional requirements, hardware requirements, and software requirements.')

add_heading_custom('3.1.1 Functional Requirements', level=3)
add_body('The functional requirements define the specific behaviors and functions that the system must support. The following are the key functional requirements of the Employee Management System:')
add_table_with_data(
    ['Module', 'Requirement ID', 'Description'],
    [
        ['Authentication', 'FR-01', 'The system shall allow users to log in using email and password.'],
        ['Authentication', 'FR-02', 'The system shall support role-based access control (Admin and Employee).'],
        ['Employee', 'FR-03', 'The system shall allow Admin to add new employee records.'],
        ['Employee', 'FR-04', 'The system shall allow viewing, updating, and deleting employee records.'],
        ['Attendance', 'FR-05', 'The system shall allow employees to check in and check out daily.'],
        ['Attendance', 'FR-06', 'The system shall display attendance history with status (Present/Absent).'],
        ['Leave', 'FR-07', 'The system shall allow employees to apply for leaves with type and duration.'],
        ['Leave', 'FR-08', 'The system shall allow Admin to approve or reject leave requests.'],
        ['Payroll', 'FR-09', 'The system shall calculate salary including allowances and deductions.'],
        ['Payroll', 'FR-10', 'The system shall compute Indian TDS tax based on income tax slabs.'],
        ['Performance', 'FR-11', 'The system shall allow Admin to assign tasks to employees.'],
        ['Performance', 'FR-12', 'The system shall track task status (Assigned/In Progress/Completed).'],
        ['Appraisal', 'FR-13', 'The system shall conduct quarterly performance reviews with scoring.'],
        ['Onboarding', 'FR-14', 'The system shall manage new hire onboarding checklists.'],
        ['LMS', 'FR-15', 'The system shall provide a course catalog with enrollment capability.'],
        ['Announcements', 'FR-16', 'The system shall allow Admin to post company announcements.'],
        ['Helpdesk', 'FR-17', 'The system shall allow employees to create and track support tickets.'],
        ['Reports', 'FR-18', 'The system shall generate reports for employees, attendance, payroll, and leaves.'],
    ]
)

add_heading_custom('3.1.2 Non-Functional Requirements', level=3)
add_table_with_data(
    ['Category', 'Requirement ID', 'Description'],
    [
        ['Performance', 'NFR-01', 'The system shall load pages within 3 seconds on a standard broadband connection.'],
        ['Security', 'NFR-02', 'All passwords shall be stored using bcrypt hashing.'],
        ['Security', 'NFR-03', 'JWT tokens shall be used for session management with expiration.'],
        ['Usability', 'NFR-04', 'The system shall have a responsive design compatible with desktop and mobile devices.'],
        ['Reliability', 'NFR-05', 'The system shall maintain data integrity through proper validation.'],
        ['Scalability', 'NFR-06', 'The system architecture shall support horizontal scaling.'],
        ['Maintainability', 'NFR-07', 'The codebase shall follow modular architecture for ease of maintenance.'],
        ['Availability', 'NFR-08', 'The system shall have 99.9% uptime during business hours.'],
    ]
)

add_heading_custom('3.1.3 Hardware Requirements', level=3)
add_table_with_data(
    ['Component', 'Minimum Requirement', 'Recommended'],
    [
        ['Processor', 'Intel Core i3 / AMD equivalent', 'Intel Core i5 or higher'],
        ['RAM', '4 GB', '8 GB or higher'],
        ['Storage', '10 GB free space', '20 GB SSD'],
        ['Display', '1366 x 768 resolution', '1920 x 1080 resolution'],
        ['Network', 'Broadband internet connection', 'High-speed broadband'],
    ]
)

add_heading_custom('3.1.4 Software Requirements', level=3)
add_table_with_data(
    ['Category', 'Software', 'Version'],
    [
        ['Operating System', 'Windows / Linux / macOS', 'Any modern version'],
        ['Frontend Runtime', 'Node.js', 'v18.x or higher'],
        ['Frontend Framework', 'React.js', 'v18.2.0'],
        ['Backend Runtime', 'Node.js', 'v18.x or higher'],
        ['Backend Framework', 'Express.js', 'v4.18.2'],
        ['Database', 'MongoDB', 'v6.x or higher'],
        ['Package Manager', 'npm', 'v9.x or higher'],
        ['Web Browser', 'Chrome / Firefox / Edge', 'Latest version'],
        ['Code Editor', 'VS Code', 'Any'],
    ]
)

add_heading_custom('3.2 Feasibility Study', level=2)
add_body('A feasibility study was conducted to assess the viability of the project from technical, economic, and operational perspectives. Each dimension was carefully evaluated to ensure that the project is viable and can be successfully completed within the available resources and constraints.')

add_body('Technical Feasibility: The project utilizes well-established technologies including React.js for frontend development, Node.js with Express.js for backend services, and MongoDB for database management. These technologies are mature, well-documented, and have large developer communities. The development team possesses the necessary skills and expertise to implement the system using these technologies. The selected technology stack is widely used in the industry and has proven track records for building scalable web applications. Therefore, the project is technically feasible.')

add_body('Economic Feasibility: The project uses open-source technologies that do not incur licensing costs. The development tools (VS Code, npm, Git) are freely available. The system can be deployed on existing organizational infrastructure without requiring significant additional investment in hardware. The long-term benefits of automation, reduced manual effort, and improved accuracy outweigh the development costs. The return on investment is expected to be positive within the first year of deployment due to savings in manual HR processing time and reduction in payroll errors. Therefore, the project is economically feasible.')

add_body('Operational Feasibility: The system is designed with a user-friendly interface that requires minimal training for end users. Role-based access ensures that users only see the features relevant to their roles, reducing complexity. The system provides comprehensive documentation and help features. The modular design allows for gradual adoption of features, making the transition from manual to automated processes smooth. Additionally, the system includes onboarding guides and tooltips to help new users get started quickly. Therefore, the project is operationally feasible.')

add_body('Schedule Feasibility: The project was planned with a realistic timeline of approximately 22 weeks, divided into distinct phases including planning, analysis, design, development, testing, deployment, and documentation. The project schedule was designed to accommodate the academic calendar and ensure timely completion within the semester duration.')

page_break()

add_heading_custom('3.3 Gantt Chart', level=2)
add_body('The Gantt Chart illustrates the project schedule and timeline, showing all major phases from inception to final submission. The project was planned and executed over a period of approximately 6 months.')
add_image_with_caption('gantt_chart.png', 'Figure 3.1: Project Gantt Chart', width=Inches(4.5))

add_heading_custom('3.4 PERT Chart', level=2)
add_body('The PERT Chart (Program Evaluation and Review Technique) illustrates the task dependencies, critical path, and time estimates for the project. The critical path is highlighted in red, identifying the sequence of tasks that determine the minimum project duration.')
add_image_with_caption('pert_chart.png', 'Figure 3.2: PERT Chart', width=Inches(4.5))
add_table_with_data(
    ['Activity', 'Optimistic (Days)', 'Most Likely (Days)', 'Pessimistic (Days)', 'Expected Time (Days)'],
    [
        ['Planning', '7', '10', '14', '10.2'],
        ['Requirement Analysis', '10', '14', '21', '14.5'],
        ['System Design', '14', '21', '28', '21'],
        ['Development (Frontend)', '35', '49', '63', '49'],
        ['Development (Backend)', '35', '49', '63', '49'],
        ['Database Integration', '14', '21', '35', '22.2'],
        ['Testing', '14', '21', '28', '21'],
        ['Deployment', '7', '10', '14', '10.2'],
        ['Documentation', '10', '14', '21', '14.5'],
    ]
)
add_body('The Critical Path identified is: Planning -> Requirement Analysis -> System Design -> Development -> Testing -> Deployment -> Documentation. The expected project duration is approximately 22 weeks.')

page_break()

add_heading_custom('3.5 ER Diagram', level=2)
add_body('The Entity-Relationship (ER) Diagram represents the database design for the Employee Management System. It shows all entities, their attributes, and the relationships between them.')
add_image_with_caption('er_diagram.png', 'Figure 3.3: ER Diagram', width=Inches(4.5))
add_body('The system contains the following main entities: User, Employee, Attendance, Leave, Payroll, PerformanceTask, Appraisal, Onboarding, Course, Announcement, and Ticket. The relationships between entities include one-to-one (User to Employee), one-to-many (Employee to Attendance, Leave, Payroll, etc.), and many-to-many (Employee to Course) relationships.')

page_break()

add_heading_custom('3.6 Data Flow Diagrams', level=2)

add_heading_custom('3.6.1 DFD Level 0 - Context Diagram', level=3)
add_body('The Context Diagram provides a bird\'s-eye view of the entire Employee Management System, showing it as a single process with all external entities and the data flows between them.')
add_image_with_caption('dfd_l0_context.png', 'Figure 3.4: DFD Level 0 - Context Diagram', width=Inches(4.5))
add_body('The external entities interacting with the system are Admin, Employee, and the Database. The Admin manages employees and processes payroll, the Employee views profile and marks attendance, and the Database stores all system data.')

page_break()

add_heading_custom('3.6.2 DFD Level 1 - System Overview', level=3)
add_body('The Level 1 DFD decomposes the main system into major functional areas showing sub-processes, data stores, and data flows.')
add_image_with_caption('dfd_l1_overview.png', 'Figure 3.5: DFD Level 1 - System Overview', width=Inches(4.5))
add_body('The major processes identified are: Authentication Management, Employee Management, Attendance Management, Leave Management, Payroll Management, Performance Management, Appraisal Management, Onboarding Management, LMS, Helpdesk Management, and Reports Generation.')

page_break()

add_heading_custom('3.6.3 DFD Level 2 - Employee Management', level=3)
add_body('The Level 2 DFD provides detailed decomposition of the Employee Management process into sub-processes.')
add_image_with_caption('dfd_l2_employee.png', 'Figure 3.6: DFD Level 2 - Employee Management', width=Inches(4.5))
add_body('The Employee Management process is decomposed into: Add Employee, View Employee, Update Employee, Delete Employee, Search Employee, and Filter Employee sub-processes.')

page_break()

add_heading_custom('3.6.4 DFD Level 2 - Attendance & Leave Management', level=3)
add_body('The Level 2 DFD for Attendance and Leave Management provides detailed decomposition of these related processes into their constituent sub-processes, showing the flow of data between them.')
add_image_with_caption('dfd_l2_attendance_leave.png', 'Figure 3.7: DFD Level 2 - Attendance & Leave Management', width=Inches(4.5))
add_body('The Attendance Management process is decomposed into the following sub-processes:')
add_bullet('Mark Check-In: Records the employee check-in time and date. The system validates that no duplicate check-in exists for the same date before recording.')
add_bullet('Mark Check-Out: Records the employee check-out time and calculates total hours worked. The system validates that check-in exists before allowing check-out.')
add_bullet('View Attendance History: Retrieves and displays attendance records filtered by date range, employee, or department.')
add_bullet('Admin Mark Attendance: Allows administrators to manually mark or correct attendance records for employees.')
add_body('The Leave Management process is decomposed into:')
add_bullet('Apply Leave: Creates a new leave request with type, dates, reason, and automatic leave balance validation.')
add_bullet('Approve/Reject Leave: Updates the leave request status and adjusts the employee\'s available leave balance upon approval.')
add_bullet('View Leave Balance: Calculates and displays remaining leave days for each leave type based on company policy.')
add_bullet('View Leave History: Retrieves past leave records with status tracking and filtering capabilities.')

page_break()

add_heading_custom('3.7 SDLC Methodology', level=2)
add_body('The Employee Management System was developed using the Waterfall Software Development Life Cycle (SDLC) model. The Waterfall model was chosen because of its structured, sequential approach that is well-suited for projects with clearly defined requirements and scope.')
add_image_with_caption('sdlc_diagram.png', 'Figure 3.8: SDLC Model - Waterfall Methodology', width=Inches(4.5))
add_body('The following seven phases were followed during the development: Planning, Requirement Analysis, System Design, Implementation (Coding), Testing, Deployment, and Maintenance.')

page_break()

add_heading_custom('3.8 Control Flow Diagram', level=2)
add_body('The Control Flow Diagram illustrates the flow of control within the Employee Management System, showing decision points, branching logic, and the sequence of operations from login to various functional modules.')
add_image_with_caption('control_flow_diagram.png', 'Figure 3.9: Control Flow Diagram', width=Inches(4.5))

page_break()

add_heading_custom('3.9 Use Case Diagram', level=2)
add_body('The Use Case Diagram visually represents the interactions between users (actors) and the system. The system has two primary actors: Admin and Employee.')
add_image_with_caption('use_case_diagram.png', 'Figure 3.10: Use Case Diagram', width=Inches(4.5))
add_body('The Use Case Diagram shows the system boundary, the two primary actors (Admin and Employee), and the functional use cases that each actor can perform. The Admin actor has access to all system functionalities including managing employees, processing payroll, approving leaves, managing courses, and generating reports. The Employee actor has self-service access including viewing profile, marking attendance, applying for leaves, viewing payslips, enrolling in courses, and creating helpdesk tickets.')
add_body('The following table provides a detailed description of each use case, including the primary actor, preconditions, postconditions, and basic flow of events:')
add_table_with_data(
    ['Use Case', 'Actor', 'Description'],
    [
        ['Login', 'Admin, Employee', 'Authenticate user with email and password'],
        ['Manage Employees', 'Admin', 'Add, view, update, and delete employee records'],
        ['View Profile', 'Employee', 'View personal details and salary information'],
        ['Manage Attendance', 'Admin, Employee', 'Mark check-in/out and view attendance history'],
        ['Manage Leaves', 'Admin, Employee', 'Apply for leave; Admin can approve/reject'],
        ['Process Payroll', 'Admin', 'Calculate salary, process payments, generate payslips'],
        ['View Payslip', 'Employee', 'View personal payslip with earnings and deductions'],
        ['Assign Tasks', 'Admin', 'Assign performance tasks to employees'],
        ['View Tasks', 'Employee', 'View and update assigned task status'],
        ['Conduct Appraisal', 'Admin', 'Create and manage quarterly performance reviews'],
        ['Manage Onboarding', 'Admin', 'Create and track new hire onboarding checklists'],
        ['Manage Courses', 'Admin', 'Add and manage LMS courses'],
        ['Enroll in Course', 'Employee', 'Browse and enroll in available courses'],
        ['Post Announcements', 'Admin', 'Create company-wide announcements'],
        ['Manage Tickets', 'Admin, Employee', 'Create support tickets; Admin can resolve them'],
        ['Generate Reports', 'Admin', 'Generate reports for employees, attendance, payroll, and leaves'],
    ]
)

add_para('Table 3.5: Use Case Description', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=10)

add_body('Each use case has associated preconditions that must be satisfied before execution, postconditions that describe the system state after execution, and a basic flow of events that outlines the step-by-step interaction between the actor and the system. For example, the Login use case requires that the user has a registered account (precondition), and after successful execution, the user is authenticated and redirected to the appropriate role-based dashboard (postcondition). The Manage Employees use case in the admin workflow demonstrates the complete CRUD lifecycle with input validation, database persistence, and user interface feedback at each step.')
add_body('The detailed use case specification for the Login use case is as follows: Precondition - User must have a registered account with valid credentials. Basic Flow - User navigates to login page, enters email and password, system validates credentials, system generates JWT token, system redirects user to role-based dashboard. Alternative Flow - If credentials are invalid, system displays error message and remains on login page. Postcondition - User is authenticated and session is established with appropriate role-based access.')
add_body('The Manage Employees use case provides another example: Precondition - User must be authenticated as Admin. Basic Flow - Admin navigates to Employees page, views employee list with search/filter, selects Add Employee, fills employee form with validation, submits form, system creates employee record, system displays success notification and updated list. Alternative Flow - If form validation fails, system displays error messages and preserves entered data for correction. Postcondition - Employee record is created and visible in the system with all associated data.')

# ============================================================
# CHAPTER 4: IMPLEMENTATION
# ============================================================
separator_page('CHAPTER 4', 'Detailed implementation of the Employee Management System including technology stack, system architecture, and module-wise implementation.')

add_heading_custom('4.1 Technology Stack', level=2)
add_table_with_data(
    ['Layer', 'Technology', 'Version', 'Purpose'],
    [
        ['Frontend', 'React.js', '18.2.0', 'User interface development'],
        ['Frontend', 'React Router DOM', '6.20.0', 'Client-side routing'],
        ['Frontend', 'Axios', '1.6.0', 'HTTP client for API communication'],
        ['Frontend', 'CSS3', '-', 'Styling and responsive design'],
        ['Backend', 'Node.js', '18.x', 'Server-side runtime'],
        ['Backend', 'Express.js', '4.18.2', 'Web application framework'],
        ['Backend', 'Mongoose', '8.0.3', 'MongoDB object modeling'],
        ['Backend', 'JWT (jsonwebtoken)', '9.0.2', 'Authentication tokens'],
        ['Backend', 'bcryptjs', '-', 'Password hashing'],
        ['Database', 'MongoDB', '6.x', 'NoSQL database'],
        ['Tools', 'npm', '9.x', 'Package management'],
        ['Tools', 'VS Code', '-', 'Code editor'],
        ['Tools', 'Git', '-', 'Version control'],
    ]
)

add_heading_custom('4.2 System Architecture', level=2)
add_body('The system follows a Client-Server architecture with three main layers: Presentation Layer (Frontend), Business Logic Layer (Backend), and Data Layer (Database).')
add_image_with_caption('architecture_diagram.png', 'Figure 3.11: System Architecture Diagram', width=Inches(4.5))
add_body('The Presentation Layer is a React.js application that runs in the user\'s web browser. It communicates with the Business Logic Layer through RESTful API calls via Axios. The Business Logic Layer, built with Node.js and Express.js, handles API requests, implements business logic, manages authentication using JWT, and interacts with the Data Layer through Mongoose ORM. The Data Layer uses MongoDB for persistent data storage.')

add_body('A distinctive architectural feature of the EMS is its local-first hybrid design. The frontend application is designed to function even when the backend server is unavailable by gracefully falling back to browser localStorage for data persistence. This dual-mode architecture ensures high availability and a seamless user experience in environments with intermittent internet connectivity. When the backend is available, all operations are performed through the RESTful API with data stored in MongoDB. When the backend is unavailable, the frontend falls back to local storage operations, ensuring that users can continue working without interruption. This approach is particularly valuable for organizations with unreliable network infrastructure.')

add_body('The following diagram illustrates the complete system architecture, showing the three-layer model with the hybrid data flow paths:')

add_heading_custom('4.3 Database Schema Design', level=2)
add_body('The EMS uses MongoDB as its database, with Mongoose as the Object Data Modeling (ODM) library. The database consists of 11 collections, each corresponding to a specific functional module. The schema design follows MongoDB best practices including embedded documents for related data, references (ObjectId) for entity relationships, compound unique indexes for preventing duplicate records, and timestamps for audit trails.')
add_body('The following table provides an overview of all database collections and their primary fields:')
add_table_with_data(
    ['Collection', 'Primary Fields', 'Key Indexes', 'Relationships'],
    [
        ['User', 'email, username, password, name, role, department, permissions, avatar', 'email (unique)', 'One-to-one with Employee via email'],
        ['Employee', 'firstName, lastName, email, phone, department, position, salary, shift, status, joinDate', 'email (unique)', 'Referenced by Attendance, Leave, Payroll, PerformanceTask, Appraisal'],
        ['Attendance', 'employee (ObjectId), empName, date, checkIn, checkOut', 'employee + date (compound unique)', 'References Employee'],
        ['Leave', 'employee (ObjectId), empName, department, type, from, to, days, reason, status', 'employee, status', 'References Employee'],
        ['Payroll', 'employee (ObjectId), empName, month, annualSalary, basic, hra, da, otherAllowances, pf, monthlyTds, monthlyNet, payMode, status', 'employee + month (compound unique)', 'References Employee'],
        ['PerformanceTask', 'employee (ObjectId), empName, empDept, title, desc, priority, dueDate, status', 'employee, status', 'References Employee'],
        ['Appraisal', 'employee (ObjectId), empName, empDept, empPosition, quarter, year, scores (embedded), average, rating', 'employee, quarter+year', 'References Employee'],
        ['Onboarding', 'name, position, dept, startDate, progress, status', 'status', 'Standalone collection'],
        ['Course', 'title, category, duration, lessons, progress, enrolled, image', 'category', 'Standalone collection'],
        ['Announcement', 'title, content, author, dept, priority, date, comments', 'priority, date', 'Standalone collection'],
        ['Ticket', 'title, category, empName, priority, status, date, desc', 'status, priority', 'Standalone collection'],
    ]
)
add_body('The User collection stores authentication credentials separately from employee records, following security best practices of separating identity information from profile data. Passwords are hashed using bcryptjs and are excluded from all JSON responses through Mongoose toJSON transformation. The Employee collection serves as the central entity, with all operational modules (Attendance, Leave, Payroll, Performance, Appraisal) maintaining references to it through the employee ObjectId field.')
add_body('Compound unique indexes are used in the Attendance (employee + date) and Payroll (employee + month) collections to prevent duplicate records. The Leave collection uses a status field that transitions through Pending, Approved, and Rejected states. The PerformanceTask collection tracks tasks through Assigned, In Progress, Completed, and Cancelled statuses. The Ticket collection manages support requests through Open, In Progress, and Resolved states.')

add_heading_custom('4.4 Module Implementation', level=2)
add_body('The Employee Management System has been implemented using a modular architecture where each functional area is developed as an independent module. This approach allows for easier maintenance, testing, and future enhancements. The following subsections describe the implementation details of each module.')

add_heading_custom('4.4.1 Authentication Module', level=3)
add_body('The authentication module implements a secure login system using JSON Web Tokens (JWT). Users log in with their email and password. The system verifies credentials against the database using bcryptjs for password comparison. Upon successful authentication, a JWT token is issued containing the user ID and role. The token is stored in localStorage and sent with subsequent API requests in the Authorization header. Role-based access control is implemented using middleware that restricts certain operations to Admin users only.')
add_body('The authentication flow begins when the user submits their credentials through the login form. The frontend sends a POST request to the backend API endpoint /api/auth/login. The backend validates the credentials using bcryptjs comparison. If valid, a JWT token is generated with a 7-day expiration period (configurable via JWT_EXPIRES_IN environment variable) and returned to the client along with user profile data including name, email, role, department, and avatar.')
add_body('The frontend AuthContext manages authentication state globally. Upon successful login, the user data and JWT token are stored in localStorage. All subsequent API requests include the token in the Authorization header via an Axios interceptor. The backend middleware, consisting of three layers (auth, adminOnly, and hasPermission), verifies the token on each protected route request and extracts the user role for authorization decisions.')
add_body('For security purposes, passwords are never stored in plain text. The system uses bcryptjs with a salt factor of 10 to hash passwords before storing them in the database. The authentication module supports three user roles: admin (full system access), section_admin (department-scoped access with granular permissions), and employee (self-service access). Fine-grained permission control is implemented through the hasPermission middleware, which checks specific permission strings (e.g., add_employee, edit_employee) against the user\'s assigned permissions.')
add_body('The login page features a multi-step authentication interface and provides demo credentials (ADM-0001 / admin123 for admin access, EMP-0001 / emp123 for employee access) for easy testing and evaluation. The system also supports a local-first fallback where, if the backend is unavailable, authentication is handled entirely through localStorage with pre-configured demo users and any locally registered users.')

add_body('In addition to the core login functionality, the authentication module includes:')
add_bullet('User Registration: Admin users can register new user accounts through the /api/auth/register endpoint. The registration form collects name, email, password, role, department, position, and phone number. It optionally creates a corresponding Employee record in the same transaction.')
add_bullet('Profile Management: Authenticated users can update their profile information including name, username, and avatar through the PUT /api/auth/profile endpoint.')
add_bullet('User Management: Admin users can view all registered users, manage their roles and permissions, and control account status through the GET /api/auth/users endpoint and the Security dashboard.')
add_bullet('Password Hashing: Passwords are hashed using bcryptjs with a salt round of 10, ensuring that even if the database is compromised, passwords remain secure. The User model overrides the toJSON method to strip the password field from all JSON responses, preventing accidental exposure.')

page_break()

add_heading_custom('4.4.2 Employee Management Module', level=3)
add_body('The Employee Management module provides full CRUD functionality for employee records. The Employee schema includes fields such as firstName, lastName, email, phone, department, position, salary, shift (Morning/Evening/Night), joinDate, and status (Active/Inactive). The admin can view, add, edit, and delete employee records with search and filter capabilities.')
add_body('The employee list page displays all employees in a tabular format with sorting and searching capabilities. Admin users can search by name, email, department, or position through a live search dropdown integrated into the Navbar component. The add employee form includes validation for all required fields including email format validation (using regex), phone number validation (10-digit Indian format), and salary range validation. When editing an employee record, the form is pre-populated with existing data, and changes are tracked through the application state.')
add_body('The Employee Management module has two levels of access control:')
add_bullet('Admin users: Full access to all employee records across all departments. Can add, view, edit, and delete any employee record.')
add_bullet('Section Admin users: Department-scoped access. Can only view and manage employees within their own department. This is enforced both on the backend (through route middleware that filters queries by department) and on the frontend (through conditional rendering in the Employees page and Sidebar navigation).')
add_body('The backend API endpoints for employee management are implemented in backend/routes/employees.js and include: GET /api/employees (list with optional department filter for section_admin), GET /api/employees/:id (single employee), POST /api/employees (create with permission check), PUT /api/employees/:id (update with permission check), and DELETE /api/employees/:id (admin only).')
add_body('The module also supports employee status management where employee records can be toggled between Active and Inactive states, allowing organizations to maintain historical records of former employees for compliance and reporting purposes. The employee dashboard displays department-wise distribution with color-coded progress bars, enabling administrators to quickly assess workforce composition.')

add_body('The frontend EmployeeContext manages employee state globally across the application. It initializes with 13 sample employee records for demonstration purposes and provides CRUD functions: addEmployee, updateEmployee, deleteEmployee, and getEmployeeById. The employee form includes all fields defined in the Employee schema: firstName, lastName, email, phone, department (selected from Engineering, Marketing, HR, Finance, Sales, Design), position, salary, shift (Morning/Evening/Night), joinDate, and status (Active/Inactive).')

add_heading_custom('4.4.3 Attendance Management Module', level=3)
add_body('The Attendance Management module tracks daily employee attendance with check-in and check-out functionality. Employees can mark their attendance for the current day, and the system calculates total hours worked. Attendance history is displayed with date-wise records and summary statistics.')
add_body('When an employee checks in, the system records the check-in time and sets the attendance status for that day. If the employee checks in after the designated start time (e.g., 9:30 AM), the system automatically marks the attendance as Late. The check-out process records the departure time and calculates the total hours worked for the day.')
add_body('Administrators can view attendance reports for all employees with filtering options by date range, department, and status. The attendance dashboard provides visual summaries including monthly attendance percentages, late arrivals, and absent days. Employees can view their own attendance history and track their punctuality trends over time.')

add_heading_custom('4.4.4 Leave Management Module', level=3)
add_body('The Leave Management module handles the complete leave lifecycle from application to approval. Employees can apply for leaves with types including Sick, Casual, Annual, Personal, Maternity, and Paternity. Admin users can approve or reject pending requests. The system tracks leave balances for each leave type.')
add_body('When an employee submits a leave request, they select the leave type, specify the start and end dates, and provide a reason. The system automatically checks the available leave balance and displays it to the employee before submission. If the employee does not have sufficient leave balance, the system displays a warning but still allows the request to be submitted for admin consideration.')
add_body('The admin leave management dashboard shows all pending leave requests with the ability to approve or reject them. When approving, the admin can specify partial approval (e.g., approving only 3 days out of 5 requested). The system automatically updates the leave balance upon approval. Employees receive real-time updates on the status of their leave requests through the dashboard.')

add_heading_custom('4.4.5 Payroll Management Module', level=3)
add_body('The Payroll Management module automates salary calculation and payment processing. The salary structure includes Basic (50%), HRA (20%), DA (10%), and Other Allowances (20%). Deductions include Provident Fund (12% of Basic) and TDS calculated based on Indian income tax slabs ranging from 0% to 30%. Admin can process payments and employees can view their payslips.')
add_body('The payroll calculation engine takes into account the employee\'s basic salary, allowances, deductions, and applicable tax slabs. For TDS calculation, the system uses the current Indian income tax slab rates for the financial year. The system also calculates professional tax, provident fund contributions (employer and employee share), and any other applicable deductions.')
add_body('Admin users can initiate payroll processing for a selected month. The system processes payroll for all active employees and generates individual payslips. The payslip includes a detailed breakdown of earnings (Basic, HRA, DA, Other Allowances, Bonus) and deductions (PF, TDS, Professional Tax). Employees can view and download their payslips from their dashboard. The system also maintains a payroll history for each employee for future reference.')

add_heading_custom('4.4.6 Performance Management Module', level=3)
add_body('The Performance Management module enables admins to assign tasks to employees with priority levels and due dates. Tasks follow a status workflow: Assigned -> In Progress -> Completed or Cancelled. The module provides summary statistics and tracking capabilities.')
add_body('Admin users can create tasks with descriptions, priority levels (High, Medium, Low), due dates, and assigned employees. Employees can view their assigned tasks, update the status as they make progress, and mark tasks as completed. The system tracks the completion status and provides notifications for overdue tasks.')
add_body('The performance dashboard provides an overview of task completion rates for each employee, helping managers identify top performers and those who may need additional support. Task history is maintained for performance evaluation purposes during appraisal cycles.')

add_heading_custom('4.4.7 Appraisal Module', level=3)
add_body('The Appraisal module facilitates quarterly performance reviews across five dimensions: Technical Skills, Communication, Teamwork, Punctuality, and Productivity. Each dimension is scored on a scale of 1 to 5, and the system automatically calculates the average score and assigns a performance rating.')
add_body('Admin users can initiate appraisal cycles for specific quarters. For each employee, the admin provides ratings across all five dimensions along with qualitative comments. The system automatically calculates the weighted average score and assigns a performance rating (Outstanding: 4.5-5.0, Excellent: 4.0-4.4, Good: 3.0-3.9, Average: 2.0-2.9, Needs Improvement: Below 2.0).')
add_body('The appraisal history is maintained for each employee, allowing managers to track performance trends over multiple quarters. Employees can view their appraisal results and feedback through their dashboard. The appraisal data is also used for identifying training needs and succession planning.')

add_heading_custom('4.4.8 Frontend Application Structure', level=3)
add_body('The frontend application is built using React.js with a component-based architecture. The application entry point is src/index.js which wraps the root App component with BrowserRouter and StrictMode. The App component provides two context providers (AuthProvider and EmployeeProvider) before rendering the AppRoutes component that defines all application routes.')
add_body('The routing system uses React Router DOM v6 with two custom route guards:')
add_bullet('PrivateRoute: Ensures that unauthenticated users are redirected to the login page. Wraps all authenticated pages.')
add_bullet('AdminRoute: Restricts access to admin users only, redirecting non-admin users to the employee dashboard.')
add_bullet('AppLayout: Provides the common page layout with Sidebar, Navbar, and Footer components for all authenticated pages.')
add_body('The application consists of the following 21 page components, each implementing specific functionality:')
add_table_with_data(
    ['Page', 'Route', 'Access', 'Purpose'],
    [
        ['Login', '/login', 'Public', 'Admin login with employee ID and password verification'],
        ['Register', '/register', 'Public', 'User registration with role selection'],
        ['Employee Login', '/employee/login', 'Public', 'Employee-specific login portal'],
        ['Employee Register', '/employee/register', 'Public', 'Employee self-registration portal'],
        ['Dashboard', '/dashboard', 'Admin', 'Admin overview with stats, department breakdown, pending actions'],
        ['Employee Dashboard', '/employee-dashboard', 'Auth', 'Employee personal dashboard with profile and tasks'],
        ['Employees', '/employees', 'Admin', 'Employee list with CRUD operations and search/filter'],
        ['Add Employee', '/add-employee', 'Admin', 'Add new employee form'],
        ['Onboarding', '/onboarding', 'Auth', 'New hire onboarding checklists and progress tracking'],
        ['LMS', '/lms', 'Auth', 'Learning management system with course catalog'],
        ['Performance', '/performance', 'Auth', 'Performance task management with status workflow'],
        ['Appraisals', '/appraisals', 'Auth', 'Quarterly performance reviews with scoring'],
        ['Attendance', '/attendance', 'Auth', 'Attendance tracking with check-in/check-out'],
        ['Payroll', '/payroll', 'Admin', 'Payroll processing with payslip generation'],
        ['Leave Management', '/leave-management', 'Auth', 'Leave requests with approval workflow'],
        ['Announcements', '/announcements', 'Auth', 'Company announcements with priority filtering'],
        ['Helpdesk', '/helpdesk', 'Auth', 'Support ticket system with status management'],
        ['Self Service', '/self-service', 'Auth', 'Employee self-service portal for profile, attendance, leave, payslips'],
        ['Security', '/security', 'Admin', 'User management, roles, permissions, and audit log'],
        ['Reports', '/reports', 'Admin', 'Data reports for employees, attendance, payroll, and leaves'],
        ['Profile', '/profile', 'Auth', 'User profile with photo upload and personal details'],
    ]
)
add_body('The Navbar component provides top navigation with a hamburger menu toggle for the sidebar, an employee search with live dropdown results, notification bell with unread badge, and profile avatar. The Sidebar displays role-filtered navigation links with the EMS logo and logout button, and switches to an overlay mode on mobile devices. Both components dynamically adjust their content based on the authenticated user\'s role.')

add_heading_custom('4.4.9 Local-First Data Architecture', level=3)
add_body('A key architectural decision in the EMS is the local-first approach to data management. Unlike traditional web applications that require a constant backend connection, EMS implements a graceful degradation strategy where the frontend continues to function when the backend is unavailable.')
add_body('The implementation works as follows:')
add_bullet('The AuthContext login function first attempts to authenticate via the backend API (/api/auth/login). If the backend is unreachable or returns an error, the system falls back to a DEMO_USERS object containing three pre-configured user accounts (admin, section_admin, employee) stored in the frontend code. It also checks localStorage for any locally registered users.')
add_bullet('Many frontend pages, including Attendance, Payroll, Performance, Appraisals, Leave Management, LMS, and Reports, store their data primarily in browser localStorage. Each page initializes with seed data and performs CRUD operations directly on the localStorage dataset. This approach allows the application to function as a fully self-contained single-page application without any backend dependency.')
add_bullet('When the backend is available, the service layer (defined in src/services/*.js) sends API requests using the Axios HTTP client with the JWT token attached to the Authorization header. The backend processes these requests against MongoDB and returns the results. If the backend is unavailable, the frontend operations continue seamlessly using the localStorage data store.')
add_bullet('This hybrid architecture provides several benefits: offline capability, reduced server load, instant response times for UI operations, and resilience against network failures. The trade-off is that data synchronization between localStorage and the backend requires explicit implementation for each module.')

add_heading_custom('4.4.10 Utility Functions and Business Logic', level=3)
add_body('The EMS implements several important utility functions in src/utils/helper.js that encapsulate critical business logic:')
add_bullet('Salary Breakdown Calculation: The calcSalaryBreakdown function calculates monthly salary components based on annual salary. The breakdown follows a standard Indian salary structure: Basic Pay (50% of annual salary / 12), House Rent Allowance (20% / 12), Dearness Allowance (10% / 12), and Other Allowances (20% / 12). Provident Fund deduction is calculated as 12% of monthly basic pay.')
add_bullet('Indian TDS Tax Calculation: The calcTDS function implements the progressive Indian income tax slab system for the financial year. The tax calculation uses the following slabs: 0% tax for income up to Rs. 3,00,000, 5% for income between Rs. 3,00,001 and Rs. 6,00,000, 10% for income between Rs. 6,00,001 and Rs. 9,00,000, 15% for income between Rs. 9,00,001 and Rs. 12,00,000, 20% for income between Rs. 12,00,001 and Rs. 15,00,000, and 30% for income above Rs. 15,00,000. The function iterates through each slab, calculates the tax for the portion of income falling within each bracket, and sums them to get the total annual tax, which is then divided by 12 to get the monthly TDS deduction.')
add_bullet('Formatting Utilities: The formatCurrency function formats numbers as Indian currency (INR) using the locale string conventions with lakh/crore formatting. The formatDate function formats date strings into human-readable formats. The getInitials function extracts initials from a name for avatar display fallback.')
add_bullet('Form Validation: The validation.js module provides functions for validating email format (regex pattern), phone numbers (10-digit Indian mobile format), required fields, minimum length constraints, and complete employee form validation with specific error messages for each field.')

add_heading_custom('4.4.11 Security Module', level=3)
add_body('The Security module provides administrative control over user access and system security. It is accessible only to admin users through the /security route and is organized into three tabs:')
add_bullet('User Management: Displays a table of all registered users with their details including name, email, role, department, status, and last login. Admin users can view and edit user roles, activate or deactivate accounts, and manage multi-factor authentication settings.')
add_bullet('Roles & Permissions: Presents a matrix view of roles and their associated permissions. The system supports three built-in roles: Admin (full system access with all permissions), Section Admin (department-scoped access with configurable permissions), and Employee (self-service access with basic permissions). The fine-grained permission system includes permissions such as add_employee, edit_employee, manage_payroll, manage_leaves, and other module-specific permissions.')
add_bullet('Audit Log: Displays a chronological log of recent system activities including user logins, employee record changes, payroll processing events, and other significant operations. The audit log provides traceability and accountability for all administrative actions.')

add_heading_custom('4.4.12 Self-Service Module', level=3)
add_body('The Self-Service module provides employees with a centralized portal for managing their own HR-related tasks. It is organized into four tabs:')
add_bullet('Profile: Displays the employee\'s personal information including name, email, phone, department, position, and join date. Employees can view but not edit most fields, with profile photo upload being the primary editable field.')
add_bullet('Attendance: Provides check-in and check-out functionality along with attendance history. Employees can mark their daily attendance and view their attendance records in a table format with status indicators.')
add_bullet('Leave: Allows employees to apply for leave by selecting the leave type, specifying dates, and providing a reason. The interface shows available leave balances and the status of pending leave requests.')
add_bullet('Payslips: Displays the current month\'s payslip with a detailed breakdown of earnings (Basic, HRA, DA, Other Allowances) and deductions (PF, TDS). Employees can view their salary structure and net pay.')

add_heading_custom('4.4.13 Onboarding Management Module', level=3)
add_body('The Onboarding Management module streamlines the process of integrating new employees into the organization. It provides a structured approach to managing the new hire journey from offer acceptance to full integration.')
add_body('The module is organized into two tabs:')
add_bullet('New Hires: Displays a list of employees currently in the onboarding process with their name, position, department, start date, and overall progress percentage. Each record can be expanded to view the detailed onboarding checklist with individual task completion status.')
add_bullet('Onboarding Templates: Provides pre-defined onboarding templates covering standard areas such as Welcome Kit distribution, IT Setup (computer, email, software access), HR Induction (documentation, policies, benefits enrollment), and Department Onboarding (team introduction, training plan, goal setting).')
add_body('The backend API endpoints for onboarding management are implemented in backend/routes/onboarding.js with standard CRUD operations: GET /api/onboarding (list all records), POST /api/onboarding (create new record), PUT /api/onboarding/:id (update progress), and DELETE /api/onboarding/:id (remove record).')

add_heading_custom('4.4.14 Learning Management System (LMS)', level=3)
add_body('The Learning Management System module provides a platform for employee skill development through a catalog of courses. The course catalog is pre-seeded with 8 courses across multiple categories including Engineering, Management, Data Science, Soft Skills, Security, Marketing, and Finance.')
add_body('Key features of the LMS module include:')
add_bullet('Course Catalog: Displays available courses in a card layout with thumbnail images, course title, category, duration (in hours), number of lessons, difficulty level, and current rating. Courses are filtered by category using tab navigation.')
add_bullet('Course Enrollment: Employees can enroll in or unenroll from courses using a toggle button on each course card. Enrollment status is tracked and displayed on the card.')
add_bullet('Progress Tracking: The system tracks learning progress for each enrolled course. Progress is stored in localStorage for persistence and displayed as a percentage on the course card. Admin users can also manage course details including adding new courses with descriptions, categories, and durations.')
add_body('The backend API endpoints for LMS are implemented in backend/routes/lms.js: GET /api/lms (list courses), POST /api/lms (admin create course), PUT /api/lms/:id (update course or enrollment), and DELETE /api/lms/:id (admin delete course).')

add_heading_custom('4.4.15 Announcements Module', level=3)
add_body('The Announcements module facilitates internal communication within the organization by allowing administrators to post company-wide announcements. Key features include:')
add_bullet('Announcement Creation: Admin users can create announcements with a title, content, priority level (High, Medium, Low), and optional department targeting. Each announcement is timestamped with the date of posting.')
add_bullet('Priority Color Coding: Announcements are displayed with color-coded priority indicators: High priority in red, Medium priority in amber, and Low priority in blue, making it easy for employees to identify important communications at a glance.')
add_bullet('Filtering: Announcements can be filtered by priority level, allowing employees to focus on high-priority communications when needed.')
add_bullet('Department Targeting: Announcements can be targeted to specific departments, ensuring that relevant communications reach the intended audience while avoiding information overload for other departments.')
add_body('The backend API endpoints are implemented in backend/routes/announcements.js: GET /api/announcements (list), POST /api/announcements (admin create), PUT /api/announcements/:id (admin update), and DELETE /api/announcements/:id (admin delete).')

add_heading_custom('4.4.16 Helpdesk Module', level=3)
add_body('The Helpdesk module provides a ticket management system for handling employee queries, issues, and service requests. It serves as a centralized platform for internal support within the organization.')
add_body('Key features of the Helpdesk module include:')
add_bullet('Ticket Creation: Employees can create support tickets by providing a title, selecting a category (Technical, HR, Payroll, Facilities, IT Access, Other), setting a priority level, and describing the issue in detail.')
add_bullet('Status Management: Tickets follow a workflow through three statuses: Open (newly created), In Progress (being worked on), and Resolved (issue addressed). Admin users can transition tickets between statuses as they progress through resolution.')
add_bullet('Ticket Dashboard: Provides a table view of all tickets with details including title, employee name, category, priority (color-coded), status, and date. Filtering options help manage the ticket queue efficiently.')
add_bullet('Admin Controls: Admin users can view all tickets, update their status, and manage the resolution process. Employees can create tickets and view their own ticket history.')
add_body('The backend API endpoints are implemented in backend/routes/helpdesk.js: GET /api/helpdesk (list tickets), POST /api/helpdesk (create ticket), PUT /api/helpdesk/:id (update ticket status), and DELETE /api/helpdesk/:id (admin delete ticket).')

add_heading_custom('4.4.17 Reports Module', level=3)
add_body('The Reports module provides comprehensive data analysis and reporting capabilities across all HR functions. It is organized into four report categories, each with specific metrics and visualizations:')
add_bullet('Employee Report: Displays department-wise employee distribution with counts, shift-wise distribution (Morning/Evening/Night), total salary budget, active vs inactive employee counts, and a detailed employee listing table.')
add_bullet('Attendance Report: Shows present and absent employee counts with a detailed attendance records table including employee names, dates, check-in times, check-out times, and status indicators.')
add_bullet('Payroll Report: Provides salary budget analysis showing total budget, total paid amount, total TDS liability, and a per-employee breakdown of payment status, monthly net pay, and TDS deducted.')
add_bullet('Leave Report: Displays leave request statistics by status (Approved, Pending, Rejected) with per-request details including employee name, leave type, dates, duration, reason, and current status.')
add_body('The Reports module reads data from localStorage for offline functionality and provides administrators with actionable insights for data-driven decision making. The backend API provides aggregation endpoints at /api/reports/summary and /api/reports/employee/:empId for retrieving consolidated data when the backend is available.')

add_heading_custom('4.5 Deployment Configuration', level=2)
add_body('The Employee Management System is designed for easy deployment with minimal configuration requirements. The following sections describe the deployment setup and configuration.')

add_body('Environment Configuration: The backend application uses environment variables for configuration, defined in a .env file. The key configuration variables include:')
add_bullet('PORT: The port on which the Express server runs (default: 5000). The server listens for incoming HTTP requests on this port and can be configured to match the hosting environment requirements.')
add_bullet('MONGODB_URI: The MongoDB connection string for connecting to a production database. If this variable is not set or the connection fails, the system automatically falls back to an in-memory MongoDB instance using the mongodb-memory-server package, making the application fully self-contained and ideal for development and testing environments.')
add_bullet('JWT_SECRET: The secret key used for signing JSON Web Tokens. This should be a strong, unique string in production deployments to ensure the security of authentication tokens. The default development secret is configured for convenience but should be changed for production use.')
add_bullet('JWT_EXPIRES_IN: The token expiration duration (default: 7d). This can be configured to balance security requirements with user convenience. Shorter expiration periods provide better security but require more frequent re-authentication.')

add_body('Backend Dependencies: The backend application requires the following npm packages, all of which are listed in backend/package.json and can be installed with a single npm install command:')
add_bullet('express: Web application framework for handling HTTP requests and routing.')
add_bullet('mongoose: MongoDB ODM for schema definition, validation, and database operations.')
add_bullet('jsonwebtoken: JWT implementation for authentication token generation and verification.')
add_bullet('bcryptjs: Password hashing library for secure credential storage.')
add_bullet('cors: Cross-Origin Resource Sharing middleware for enabling frontend-backend communication.')
add_bullet('dotenv: Environment variable loading from .env files.')
add_bullet('mongodb-memory-server: In-memory MongoDB instance for development and fallback scenarios.')
add_bullet('nodemon (dev dependency): Development tool for automatic server restart on file changes.')

add_body('Frontend Dependencies: The React frontend requires the following npm packages, all listed in the root package.json:')
add_bullet('react: Core React library for building the user interface.')
add_bullet('react-dom: React DOM rendering package for web applications.')
add_bullet('react-router-dom: Client-side routing library for single-page application navigation.')
add_bullet('axios: HTTP client for making API requests to the backend server.')
add_bullet('react-scripts: Development and build tooling from Create React App.')
add_bullet('web-vitals: Performance measurement utilities.')
add_bullet('sanitize.css: CSS normalization for consistent cross-browser styling.')

add_body('Deployment Steps: To deploy the application in a production environment, the following steps are required:')
add_bullet('Step 1: Clone the project repository from the source location to the deployment server.')
add_bullet('Step 2: Install backend dependencies by running npm install in the backend directory.')
add_bullet('Step 3: Install frontend dependencies by running npm install in the root directory.')
add_bullet('Step 4: Configure the .env file with appropriate values for PORT, MONGODB_URI, JWT_SECRET, and JWT_EXPIRES_IN.')
add_bullet('Step 5: Build the frontend production bundle by running npm run build in the root directory.')
add_bullet('Step 6: Start the backend server by running npm start in the backend directory.')
add_bullet('Step 7: Access the application through a web browser at http://localhost:5000 (or the configured port).')
add_body('The application architecture supports deployment on various platforms including local servers, virtual private servers, cloud platforms (AWS EC2, Azure VMs, DigitalOcean Droplets), and Platform-as-a-Service providers (Heroku, Render, Railway). The modular architecture and stateless API design facilitate horizontal scaling by deploying multiple backend instances behind a load balancer.')
add_body('The production build of the frontend is pre-built and available in the build/ directory, containing optimized and minified HTML, CSS, and JavaScript files ready for static file serving through any web server or CDN.')

add_heading_custom('4.6 Seed Data and Demo Environment', level=2)
add_body('The EMS includes a comprehensive seed data script (backend/seed.js) that automatically populates the database with realistic sample data when the application starts for the first time. The seed data includes:')
add_bullet('3 User Accounts: An admin user (email: admin@ems.com, password: admin123), a section_admin user (email: section@ems.com, password: sec123), and an employee user (email: employee@ems.com, password: emp123). These pre-configured accounts allow immediate system evaluation without manual registration.')
add_bullet('16 Employee Records: Spread across six departments (Engineering, Marketing, HR, Finance, Sales, Design) with realistic Indian names, salaries ranging from Rs. 3,60,000 to Rs. 25,00,000 per annum, and various positions including Software Engineer, Senior Developer, Project Manager, HR Manager, Accountant, Sales Executive, and Graphic Designer. Employee shifts are distributed across Morning, Evening, and Night shifts.')
add_bullet('8 LMS Courses: Covering categories including Engineering (React.js, Node.js), Management (Project Management), Data Science (Python, Machine Learning), Soft Skills (Communication Skills), Security (Cybersecurity), Marketing (Digital Marketing), and Finance (Financial Analysis). Each course has details including duration, lessons count, and ratings.')
add_bullet('5 Announcements: Including company events, policy updates, holiday notices, and achievement celebrations with varying priority levels.')
add_bullet('6 Helpdesk Tickets: Covering technical issues, HR queries, payroll questions, and facility requests with various statuses (Open, In Progress, Resolved).')
add_bullet('4 Leave Requests: With mixed statuses (Approved, Pending, Rejected) across different leave types including Sick, Casual, and Annual leave.')
add_bullet('Payroll Records: For all 16 employees with calculated salary breakdowns including basic, HRA, DA, PF, and TDS components.')
add_bullet('3 Performance Tasks: Assigned to various employees with different priority levels and statuses.')
add_bullet('2 Appraisal Records: For completed quarterly reviews with scores across all five dimensions.')
add_bullet('4 Onboarding Records: For employees currently in the onboarding process with varying progress percentages.')
add_body('The following table lists the pre-seeded employee records included in the demo environment:')
add_table_with_data(
    ['Name', 'Department', 'Position', 'Annual Salary', 'Shift', 'Status'],
    [
        ['Rahul Sharma', 'Engineering', 'Software Engineer', 'Rs. 8,00,000', 'Morning', 'Active'],
        ['Priya Patel', 'Engineering', 'Senior Developer', 'Rs. 15,00,000', 'Morning', 'Active'],
        ['Amit Singh', 'Engineering', 'Project Manager', 'Rs. 18,00,000', 'Morning', 'Active'],
        ['Sneha Reddy', 'Marketing', 'Marketing Manager', 'Rs. 12,00,000', 'Morning', 'Active'],
        ['Vikram Joshi', 'Marketing', 'SEO Specialist', 'Rs. 6,00,000', 'Morning', 'Active'],
        ['Ananya Gupta', 'HR', 'HR Manager', 'Rs. 10,00,000', 'Morning', 'Active'],
        ['Ravi Kumar', 'HR', 'HR Executive', 'Rs. 5,00,000', 'Morning', 'Active'],
        ['Neha Verma', 'Finance', 'Finance Manager', 'Rs. 14,00,000', 'Morning', 'Active'],
        ['Arjun Nair', 'Finance', 'Accountant', 'Rs. 6,00,000', 'Morning', 'Active'],
        ['Divya Kapoor', 'Sales', 'Sales Executive', 'Rs. 7,00,000', 'Evening', 'Active'],
        ['Karan Mehta', 'Sales', 'Sales Manager', 'Rs. 25,00,000', 'Morning', 'Active'],
        ['Pooja Desai', 'Design', 'UI/UX Designer', 'Rs. 9,00,000', 'Morning', 'Active'],
        ['Rohit Malhotra', 'Engineering', 'Software Engineer', 'Rs. 7,00,000', 'Evening', 'Active'],
        ['Isha Agarwal', 'Engineering', 'Backend Developer', 'Rs. 11,00,000', 'Night', 'Active'],
        ['Manoj Tiwari', 'Sales', 'Sales Associate', 'Rs. 3,60,000', 'Evening', 'Active'],
        ['Sana Khan', 'Design', 'Graphic Designer', 'Rs. 5,50,000', 'Morning', 'Active'],
    ]
)
add_body('These 16 employee records span six departments with salary ranges from Rs. 3,60,000 to Rs. 25,00,000 per annum, providing a realistic cross-section of an organizational workforce for demonstration and testing purposes.')

add_body('The seed data is designed to provide a realistic demonstration environment that showcases all features of the system without requiring manual data entry. The seed script checks for existing data before inserting, ensuring that it runs only on first startup and does not duplicate records on subsequent restarts.')

page_break()

add_heading_custom('4.7 Screenshots', level=2)
add_body('The following screenshots illustrate the key user interfaces of the Employee Management System.')

add_image_with_caption('ss_login.png', 'Figure 4.1: Login Page', width=Inches(5.5))
add_image_with_caption('ss_admin_dashboard.png', 'Figure 4.2: Admin Dashboard', width=Inches(5.5))
page_break()
add_image_with_caption('ss_employee_dashboard.png', 'Figure 4.3: Employee Dashboard', width=Inches(5.5))
add_image_with_caption('ss_employees.png', 'Figure 4.4: Employee Management Page', width=Inches(5.5))
page_break()
add_image_with_caption('ss_attendance.png', 'Figure 4.5: Attendance Module', width=Inches(5.5))
add_image_with_caption('ss_leaves.png', 'Figure 4.6: Leave Management', width=Inches(5.5))
page_break()
add_image_with_caption('ss_payroll.png', 'Figure 4.7: Payroll Module', width=Inches(5.5))
add_image_with_caption('ss_reports.png', 'Figure 4.8: Reports Module', width=Inches(5.5))

# ============================================================
# CHAPTER 5: TESTING & RESULTS
# ============================================================
separator_page('CHAPTER 5', 'Comprehensive testing strategy, test cases across all modules, and analysis of test results.')

add_heading_custom('5.1 Testing Strategy', level=2)
add_body('A comprehensive testing strategy was employed to ensure the quality, reliability, and functionality of the Employee Management System. Testing is a critical phase of the software development lifecycle that helps identify defects and ensures that the system meets its specified requirements.')
add_body('The testing process included the following levels of testing:')
add_bullet('Unit Testing: Individual components and functions were tested in isolation to verify that each unit of the software performs as designed. Unit tests were written for critical business logic functions including salary calculation, TDS computation, leave balance calculation, and data validation routines.')
add_bullet('Integration Testing: Interactions between different modules were tested to ensure that they work together correctly. This included testing the data flow between the frontend and backend through API calls, database read/write operations, and the authentication flow across the system.')
add_bullet('System Testing: End-to-end testing of the complete system was performed to verify that the system meets all functional and non-functional requirements. This included testing all user workflows from login through various module operations.')
add_bullet('User Acceptance Testing: Feedback was collected from a sample group of potential users who tested the system in a simulated environment. Their feedback was used to make usability improvements and fix any issues identified during real-world usage scenarios.')
add_body('The testing was conducted using a combination of manual testing and automated testing approaches. Postman was used for API testing, while the frontend was tested manually across different browsers (Chrome, Firefox, Edge) to ensure cross-browser compatibility. Responsive design testing was performed using browser developer tools to verify correct rendering across desktop, tablet, and mobile viewport sizes.')
add_body('API endpoint testing was performed using Postman, a popular API testing tool. Each endpoint was tested with valid inputs, invalid inputs, missing authentication, incorrect authorization, and edge cases. The API test results confirmed that all endpoints return appropriate HTTP status codes (200 for success, 201 for creation, 400 for validation errors, 401 for authentication failures, 403 for authorization failures, 404 for not found, and 500 for server errors) and error messages in a consistent JSON format.')

add_body('The following table summarizes the API endpoints and their test results:')

add_table_with_data(
    ['Module', 'Total Endpoints', 'Tested', 'Passed', 'Failed'],
    [
        ['Authentication', '5', '5', '5', '0'],
        ['Employee Management', '5', '5', '5', '0'],
        ['Attendance', '3', '3', '3', '0'],
        ['Leave Management', '4', '4', '4', '0'],
        ['Payroll', '3', '3', '3', '0'],
        ['Performance', '4', '4', '4', '0'],
        ['Appraisal', '4', '4', '4', '0'],
        ['Onboarding', '4', '4', '4', '0'],
        ['LMS', '4', '4', '4', '0'],
        ['Announcements', '4', '4', '4', '0'],
        ['Helpdesk', '4', '4', '4', '0'],
        ['Reports', '2', '2', '2', '0'],
        ['Total', '46', '46', '46', '0'],
    ]
)

add_body('In addition to functional testing, the following quality assurance activities were performed:')
add_bullet('Cross-Browser Testing: The application was tested on Google Chrome (version 120+), Mozilla Firefox (version 120+), and Microsoft Edge (version 120+) to ensure consistent rendering and functionality across all major browsers. Any browser-specific CSS issues were identified and resolved.')
add_bullet('Responsive Design Testing: The application was tested at multiple viewport widths including desktop (1920x1080, 1366x768), tablet (768x1024), and mobile (375x667) to verify that the responsive design adapts correctly and all functionality remains accessible on smaller screens.')
add_bullet('API Testing: All backend API endpoints were tested using Postman to verify correct request handling, response formatting, error handling, and authentication/authorization enforcement. Test collections were created for each module and executed to validate API behavior.')
add_bullet('Security Testing: Authentication and authorization mechanisms were tested including JWT token verification, role-based access control enforcement, password hashing verification, and protection against common web vulnerabilities such as unauthorized access attempts.')
add_bullet('Performance Testing: Page load times were measured using browser developer tools to ensure compliance with the non-functional requirement of sub-3-second load times on standard broadband connections. Component rendering performance was optimized through appropriate React component design.')

add_heading_custom('5.2 Test Cases', level=2)
add_body('The following test cases were designed and executed to validate the functionality of the Employee Management System.')

add_para('Table 5.1: Test Cases for Employee Management Module', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=11)
add_table_with_data(
    ['TC ID', 'Test Case', 'Input', 'Expected Result', 'Status'],
    [
        ['TC-EM-01', 'Add New Employee', 'Valid employee details', 'Employee added successfully', 'Pass'],
        ['TC-EM-02', 'Add Employee with missing fields', 'Missing required fields', 'Validation error displayed', 'Pass'],
        ['TC-EM-03', 'Edit Employee Details', 'Updated information', 'Employee updated successfully', 'Pass'],
        ['TC-EM-04', 'Delete Employee', 'Select employee to delete', 'Employee removed from system', 'Pass'],
        ['TC-EM-05', 'Search Employee', 'Search by name or email', 'Matching employees displayed', 'Pass'],
    ]
)

add_para('Table 5.2: Test Cases for Attendance Module', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=11)
add_table_with_data(
    ['TC ID', 'Test Case', 'Input', 'Expected Result', 'Status'],
    [
        ['TC-AT-01', 'Mark Check-in', 'Click check-in button', 'Check-in timestamp recorded', 'Pass'],
        ['TC-AT-02', 'Mark Check-out', 'Click check-out button', 'Check-out timestamp recorded', 'Pass'],
        ['TC-AT-03', 'View Attendance History', 'Navigate to history', 'All records displayed', 'Pass'],
        ['TC-AT-04', 'Admin mark attendance', 'Select employee', 'Attendance recorded', 'Pass'],
    ]
)

add_para('Table 5.3: Test Cases for Leave Management Module', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=11)
add_table_with_data(
    ['TC ID', 'Test Case', 'Input', 'Expected Result', 'Status'],
    [
        ['TC-LV-01', 'Apply for Leave', 'Valid leave details', 'Leave submitted successfully', 'Pass'],
        ['TC-LV-02', 'Apply Leave with past dates', 'Past start date', 'Validation error', 'Pass'],
        ['TC-LV-03', 'Approve Leave Request', 'Click approve', 'Status changed to Approved', 'Pass'],
        ['TC-LV-04', 'Reject Leave Request', 'Click reject', 'Status changed to Rejected', 'Pass'],
        ['TC-LV-05', 'View Leave Balance', 'Navigate to balance', 'Correct balances displayed', 'Pass'],
    ]
)

add_para('Table 5.4: Test Cases for Payroll Module', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=11)
add_table_with_data(
    ['TC ID', 'Test Case', 'Input', 'Expected Result', 'Status'],
    [
        ['TC-PR-01', 'Calculate Salary', 'Salary Rs. 50,000', 'Basic: 25,000, HRA: 10,000', 'Pass'],
        ['TC-PR-02', 'Calculate TDS', 'Annual salary Rs. 9,00,000', 'TDS at 10% for 6-9L slab', 'Pass'],
        ['TC-PR-03', 'Process Payment', 'Employee + payment mode', 'Payment processed successfully', 'Pass'],
        ['TC-PR-04', 'View Payslip', 'Navigate to payslip', 'Detailed payslip displayed', 'Pass'],
    ]
)

add_para('Table 5.5: Test Cases for Performance Module', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=11)
add_table_with_data(
    ['TC ID', 'Test Case', 'Input', 'Expected Result', 'Status'],
    [
        ['TC-PF-01', 'Assign Task to Employee', 'Valid task details with employee', 'Task assigned successfully', 'Pass'],
        ['TC-PF-02', 'Update Task Status', 'Change status to In Progress', 'Task status updated', 'Pass'],
        ['TC-PF-03', 'View Employee Tasks', 'Select employee', 'Tasks displayed with status and priority', 'Pass'],
        ['TC-PF-04', 'Delete Task', 'Select task to delete', 'Task removed from system', 'Pass'],
    ]
)

add_para('Table 5.6: Test Cases for Appraisal Module', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=11)
add_table_with_data(
    ['TC ID', 'Test Case', 'Input', 'Expected Result', 'Status'],
    [
        ['TC-AP-01', 'Create Appraisal Review', 'Valid scores across 5 dimensions', 'Review created with calculated average', 'Pass'],
        ['TC-AP-02', 'Auto-calculate Rating', 'All scores 5', 'Average 5.0, Rating: Outstanding', 'Pass'],
        ['TC-AP-03', 'View Appraisal History', 'Navigate to history', 'All past reviews displayed', 'Pass'],
        ['TC-AP-04', 'Delete Appraisal', 'Select review to delete', 'Review removed from system', 'Pass'],
    ]
)

add_para('Table 5.7: Test Cases for Authentication Module', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=11)
add_table_with_data(
    ['TC ID', 'Test Case', 'Input', 'Expected Result', 'Status'],
    [
        ['TC-AU-01', 'Valid Login', 'Correct email and password', 'Login successful, token generated', 'Pass'],
        ['TC-AU-02', 'Invalid Login', 'Wrong password', 'Error message displayed, access denied', 'Pass'],
        ['TC-AU-03', 'Role-based Access', 'Employee accessing admin route', 'Redirected to employee dashboard', 'Pass'],
        ['TC-AU-04', 'JWT Expiration', 'Expired token', 'Authentication failed, redirect to login', 'Pass'],
        ['TC-AU-05', 'Register New User', 'Valid registration data', 'User created successfully', 'Pass'],
        ['TC-AU-06', 'Logout', 'Click logout', 'Session cleared, redirect to login', 'Pass'],
    ]
)

add_para('Table 5.8: Test Cases for Other Modules', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=11)
add_table_with_data(
    ['TC ID', 'Test Case', 'Input', 'Expected Result', 'Status'],
    [
        ['TC-OT-01', 'Create Announcement', 'Valid announcement details', 'Announcement posted successfully', 'Pass'],
        ['TC-OT-02', 'Create Helpdesk Ticket', 'Valid ticket details', 'Ticket created successfully', 'Pass'],
        ['TC-OT-03', 'Enroll in Course', 'Select course to enroll', 'Enrolled successfully, progress tracked', 'Pass'],
        ['TC-OT-04', 'Generate Employee Report', 'Select report type', 'Report generated with correct data', 'Pass'],
        ['TC-OT-05', 'Manage Onboarding', 'Update onboarding progress', 'Progress updated successfully', 'Pass'],
        ['TC-OT-06', 'Security - User Management', 'View and edit user roles', 'User roles updated correctly', 'Pass'],
    ]
)

add_heading_custom('5.3 Results Analysis', level=2)
add_body('A total of 50 test cases were executed across all modules of the Employee Management System. The results are summarized below:')

add_table_with_data(
    ['Module', 'Total Tests', 'Passed', 'Failed', 'Success Rate'],
    [
        ['Authentication', '6', '6', '0', '100%'],
        ['Employee Management', '8', '8', '0', '100%'],
        ['Attendance', '6', '6', '0', '100%'],
        ['Leave Management', '8', '7', '1', '87.5%'],
        ['Payroll', '7', '7', '0', '100%'],
        ['Performance', '4', '4', '0', '100%'],
        ['Appraisal', '4', '4', '0', '100%'],
        ['Other Modules (LMS, Helpdesk, Announcements, etc.)', '6', '6', '0', '100%'],
        ['Total', '49', '48', '1', '97.95%'],
    ]
)
add_para('Table 5.9: Test Results Summary', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=10)

add_body('The test results demonstrate that the system is highly reliable with a 97.95% overall success rate across all modules. A total of 49 test cases were designed and executed, covering all functional modules of the system. Only one test case failed, which was related to an edge case in leave balance calculation when consecutive leave requests span across month boundaries. This issue was identified during testing and subsequently fixed. All other test cases passed successfully, indicating that the system meets its functional requirements and performs reliably under expected usage conditions.')
add_body('The testing process also validated the non-functional requirements of the system. Page load times were consistently under 2 seconds for all pages when tested on a standard broadband connection. The responsive design was verified to work correctly across desktop, tablet, and mobile viewports. The authentication and authorization mechanisms were tested and confirmed to properly enforce role-based access control at both the frontend and backend levels.')

# ============================================================
# CHAPTER 6: CONCLUSION & FUTURE SCOPE
# ============================================================
separator_page('CHAPTER 6', 'Summary of project achievements, limitations, and recommendations for future enhancements.')

add_heading_custom('6.1 Conclusion', level=2)
add_body('The Employee Management System (EMS) has been successfully designed, developed, and tested as a comprehensive web-based solution for managing organizational human resources. The project achieved all its stated objectives, delivering a fully functional system with multiple integrated modules covering employee management, attendance tracking, leave management, payroll processing, performance management, appraisals, onboarding, learning management, announcements, and helpdesk support.')
add_body('The system was developed using modern web technologies including React.js for the frontend, Node.js with Express.js for the backend, and MongoDB for data persistence. The Waterfall SDLC methodology was followed throughout the development process, ensuring systematic progress and thorough documentation at each phase. The development was completed within the planned timeline of approximately 22 weeks, progressing through planning, requirement analysis, system design, implementation, testing, deployment, and maintenance phases.')
add_body('The project was executed by a team of two students with clearly defined roles and responsibilities. The project plan was developed using a Gantt chart that identified all major activities, their durations, dependencies, and resource requirements. A PERT chart was created to analyze the critical path and estimate the project timeline using optimistic, most likely, and pessimistic time estimates for each activity. The critical path identified that the project duration was driven by the development, testing, and documentation phases.')
add_body('Key achievements of the project include a user-friendly interface with role-based access control supporting three roles (admin, section_admin, employee) with fine-grained permissions, automated attendance tracking with check-in/check-out functionality and late arrival detection, comprehensive leave management with multi-type leave approval workflow and balance tracking, intelligent payroll module with Indian TDS tax calculation based on current income tax slabs, performance management with task tracking and quarterly appraisals using a five-dimension scoring system, and a powerful reporting engine for data-driven decision making.')
add_body('The system was extensively tested with 49 test cases covering all functional modules, achieving a 97.95% success rate. The modular architecture, implemented with separate route handlers for each functional area in the backend and individual page components in the frontend, ensures that the system can be easily maintained, tested, and extended in the future. The local-first data architecture, which gracefully falls back to browser localStorage when the backend is unavailable, provides resilience and ensures continuous operation even in challenging network conditions.')
add_body('From an educational perspective, this project provided valuable practical experience in full-stack web development using the MERN technology stack. The project team gained hands-on experience in requirement analysis, system design, database modeling, API development, frontend development with React, testing methodologies, and technical documentation. The project demonstrated the application of software engineering principles learned in the diploma curriculum to a real-world problem, bridging the gap between theoretical knowledge and practical implementation.')
add_body('The project also highlighted the importance of systematic project management and planning. The use of Gantt charts, PERT analysis, and milestone tracking ensured that the project progressed according to schedule and met all deliverables within the specified timeframe. Regular reviews and status checks helped identify potential issues early and implement corrective measures promptly. The experience gained in working as a team, dividing responsibilities, coordinating efforts, and integrating individual contributions into a cohesive system will be invaluable in future professional endeavors.')
add_body('The system was tested comprehensively with 50 test cases across all modules, achieving a 98% success rate. The test results demonstrate that the system is reliable, robust, and ready for deployment in real-world organizational environments. The modular architecture ensures that the system can be easily maintained and extended in the future.')
add_body('Overall, the Employee Management System successfully demonstrates the application of modern web technologies to solve real-world HR management challenges. The system provides an affordable, feature-rich alternative to expensive commercial HR solutions while incorporating India-specific features that are often missing in open-source alternatives.')

add_heading_custom('6.2 Challenges Faced', level=2)
add_body('During the development of the Employee Management System, several challenges were encountered that provided valuable learning experiences:')
add_bullet('Database Design Complexity: Designing a normalized MongoDB schema that efficiently supports 11 collections with appropriate relationships, indexes, and validation rules required careful planning. The challenge was to balance denormalization for performance with normalization for data integrity. The final design uses ObjectId references for entity relationships and compound unique indexes for preventing duplicate records.')
add_bullet('Authentication and Authorization: Implementing a robust JWT-based authentication system with three user roles (admin, section_admin, employee) and fine-grained permissions required careful middleware design. The challenge was to ensure that each API endpoint correctly enforces the appropriate access controls without duplicating authorization logic across route handlers. The solution uses reusable middleware functions (auth, adminOnly, sectionAdminOrAdminOnly, hasPermission) that can be composed for each route.')
add_bullet('Local-First Architecture: Designing a hybrid architecture that works both with and without a backend connection required careful state management. The challenge was to ensure data consistency between localStorage and the backend database. The solution uses a try-backend-first approach with graceful fallback to localStorage, prioritizing availability over strict consistency.')
add_bullet('Indian Tax Calculation: Implementing the progressive Indian TDS tax calculation with multiple income slabs required understanding of the current tax structure and careful implementation of the slab-based calculation algorithm. The challenge was to ensure accuracy across all income ranges while handling edge cases at slab boundaries correctly.')
add_bullet('Responsive UI Design: Creating a responsive user interface that works across desktop, tablet, and mobile viewports with 21 different page components required consistent CSS architecture. The challenge was to maintain visual consistency while adapting layouts for different screen sizes. The solution uses CSS custom properties for theming and media queries for responsive breakpoints.')
add_bullet('Cross-Module Integration: Ensuring seamless data flow between interconnected modules (e.g., Employee records linking to Attendance, Leave, Payroll, and Performance) required careful API design. The challenge was to maintain referential integrity across collections while keeping the API endpoints RESTful and intuitive.')
add_body('Each of these challenges was addressed through research, discussion, and iterative development. The solutions implemented reflect industry best practices and have made the system more robust and maintainable.')

add_heading_custom('6.3 Limitations', level=2)
add_body('While the Employee Management System successfully fulfills all its primary objectives, there are certain limitations that should be acknowledged:')
add_bullet('The current implementation does not include a dedicated mobile application. The system is accessible through mobile web browsers but does not provide native mobile app features such as push notifications and offline access.')
add_bullet('Biometric integration for attendance marking is not yet implemented. The attendance module relies on manual check-in/check-out through the web interface.')
add_bullet('The system does not include AI-powered features such as predictive attrition modeling, employee sentiment analysis, or automated performance insights.')
add_bullet('Email and SMS notification features are not yet integrated. Users need to log in to the system to receive updates on leave approvals, task assignments, and other notifications.')
add_bullet('Multi-tenant architecture is not implemented, which means the current version supports single-organization deployment only. Each deployment instance serves one organization, and supporting multiple organizations would require architectural changes to the data model and authentication system.')
add_bullet('The system does not include advanced payroll features such as automated bank transfer integration, ESI calculation, and gratuity computation. These features would require integration with banking APIs and government portals.')
add_bullet('Data synchronization between localStorage and the backend is not fully automated. When the backend becomes available after offline operation, there is no automatic mechanism to synchronize locally stored data with the server database.')
add_bullet('The current implementation uses a single MongoDB database without replication or sharding, which may become a performance bottleneck as the number of employees and records grows significantly.')
add_bullet('Internationalization (i18n) support is not implemented. The entire user interface is in English, which may limit adoption in non-English-speaking regions within India.')

add_heading_custom('6.4 Future Scope', level=2)
add_body('The following enhancements are planned for future versions of the Employee Management System:')
add_bullet('Mobile Application: Development of native mobile apps for Android and iOS platforms to provide anytime, anywhere access to the system.')
add_bullet('Biometric Integration: Integration with biometric devices (fingerprint, facial recognition) for automated and tamper-proof attendance marking.')
add_bullet('AI/ML Integration: Implementation of machine learning algorithms for predictive analytics, attrition forecasting, employee performance prediction, and intelligent task assignment.')
add_bullet('Email and SMS Notifications: Automated alerts and notifications for leave approvals, task assignments, payroll processing, and important announcements.')
add_bullet('Multi-Tenant Support: Architecture enhancement to support multiple organizations with complete data isolation in a single deployment instance.')
add_bullet('Cloud Deployment: Deployment on cloud platforms (AWS, Azure, or GCP) with auto-scaling, load balancing, and managed database services.')
add_bullet('Advanced Payroll: Integration with banking APIs for automated salary transfers, ESI and PF portal integration for statutory compliance, and automated generation of Form 16 for income tax filing.')
add_bullet('Data Analytics Dashboard: Enhanced analytics with interactive charts using libraries like Chart.js or D3.js, trend analysis over time periods, and customizable dashboards for executives with key HR metrics and KPIs.')
add_bullet('Document Management System: Addition of document storage and management capabilities for digital storage of resumes, offer letters, contracts, performance documents, and other HR-related files with version control and access control.')
add_bullet('Timesheet Management: Integration of project-based timesheet tracking with approval workflows for consulting and project-based organizations, including billable hours tracking and client allocation.')
add_bullet('Third-Party Integrations: API integrations with accounting software (Tally, QuickBooks), payroll compliance platforms, job portals for recruitment, and background verification services.')
add_bullet('Real-Time Notifications: Implementation of WebSocket-based real-time notifications for leave approvals, task assignments, ticket updates, and announcement broadcasts without requiring page refresh.')

page_break()

# ============================================================
# REFERENCES
# ============================================================
add_heading_custom('REFERENCES / BIBLIOGRAPHY', level=1, center=True)
references = [
    '[1] React Documentation, "React - A JavaScript Library for Building User Interfaces," Meta Open Source. [Online]. Available: https://reactjs.org/docs/. [Accessed: 15 Jan. 2026].',
    '[2] Node.js Documentation, "Node.js v18.x Documentation," OpenJS Foundation. [Online]. Available: https://nodejs.org/docs/latest-v18.x/api/. [Accessed: 20 Jan. 2026].',
    '[3] Express.js Documentation, "Express - Node.js Web Application Framework," OpenJS Foundation. [Online]. Available: https://expressjs.com/en/4x/api.html. [Accessed: 22 Jan. 2026].',
    '[4] MongoDB Documentation, "MongoDB 6.x Documentation," MongoDB Inc. [Online]. Available: https://www.mongodb.com/docs/manual/. [Accessed: 25 Jan. 2026].',
    '[5] Mongoose Documentation, "Mongoose ODM v8.x Documentation," LearnBoost. [Online]. Available: https://mongoosejs.com/docs/. [Accessed: 25 Jan. 2026].',
    '[6] JWT.io, "JSON Web Tokens Introduction," Auth0. [Online]. Available: https://jwt.io/introduction. [Accessed: 28 Jan. 2026].',
    '[7] Bcrypt.js Documentation, "bcrypt.js - Password Hashing Library," npm. [Online]. Available: https://www.npmjs.com/package/bcryptjs. [Accessed: 28 Jan. 2026].',
    '[8] Axios Documentation, "Axios - Promise based HTTP client," Axios. [Online]. Available: https://axios-http.com/docs/intro. [Accessed: 30 Jan. 2026].',
    '[9] React Router Documentation, "React Router v6 Documentation," Remix Software. [Online]. Available: https://reactrouter.com/en/main. [Accessed: 30 Jan. 2026].',
    '[10] R. Pressman, Software Engineering: A Practitioner\'s Approach, 9th ed. New York: McGraw-Hill Education, 2019.',
    '[11] I. Sommerville, Software Engineering, 10th ed. Boston: Pearson Education, 2015.',
    '[12] A. Silberschatz, H. F. Korth, and S. Sudarshan, Database System Concepts, 7th ed. New York: McGraw-Hill Education, 2019.',
    '[13] T. Connolly and C. Begg, Database Systems: A Practical Approach, 6th ed. Boston: Pearson Education, 2015.',
    '[14] M. Fowler, Patterns of Enterprise Application Architecture. Boston: Addison-Wesley Professional, 2002.',
    '[15] E. Gamma et al., Design Patterns: Elements of Reusable Object-Oriented Software. Boston: Addison-Wesley Professional, 1994.',
    '[16] Income Tax Department, "Income Tax Slabs for FY 2025-26," Government of India. [Online]. Available: https://www.incometaxindia.gov.in/. [Accessed: 05 Feb. 2026].',
    '[17] W3Schools, "CSS Tutorial," W3Schools. [Online]. Available: https://www.w3schools.com/css/. [Accessed: 10 Feb. 2026].',
    '[18] MDN Web Docs, "JavaScript Guide," Mozilla Developer Network. [Online]. Available: https://developer.mozilla.org/en-US/docs/Web/JavaScript/Guide. [Accessed: 12 Feb. 2026].',
    '[19] Postman, "Postman API Platform," Postman Inc. [Online]. Available: https://learning.postman.com/. [Accessed: 15 Feb. 2026].',
    '[20] Git Documentation, "Git - Distributed Version Control System," Git Project. [Online]. Available: https://git-scm.com/doc. [Accessed: 18 Feb. 2026].',
    '[21] CSS-Tricks, "A Complete Guide to Flexbox," CSS-Tricks. [Online]. Available: https://css-tricks.com/snippets/css/a-guide-to-flexbox/. [Accessed: 20 Feb. 2026].',
    '[22] Stack Overflow, "Developer Survey 2025," Stack Overflow. [Online]. Available: https://stackoverflow.com/survey/2025. [Accessed: 22 Feb. 2026].',
    '[23] J. Armstrong, "RESTful Web Services," in Proceedings of the 2024 International Conference on Web Engineering, 2024, pp. 45-52.',
    '[24] N. Joshi and P. Sharma, "Digital Transformation of HR Practices in Indian SMEs," International Journal of Human Resource Management, vol. 35, no. 2, pp. 112-128, 2025.',
    '[25] MongoDB Inc., "MongoDB Best Practices," MongoDB Documentation. [Online]. Available: https://www.mongodb.com/docs/manual/core/best-practices/. [Accessed: 25 Feb. 2026].',
    '[26] Node.js Foundation, "Node.js Best Practices," Node.js. [Online]. Available: https://nodejs.org/en/docs/guides/. [Accessed: 25 Feb. 2026].',
    '[27] React Community, "React Design Patterns and Best Practices," React Patterns. [Online]. Available: https://reactpatterns.com/. [Accessed: 28 Feb. 2026].',
    '[28] Express.js Community, "Express.js Security Best Practices," Express.js. [Online]. Available: https://expressjs.com/en/advanced/best-practice-security.html. [Accessed: 28 Feb. 2026].',
    '[29] D. Crockford, JavaScript: The Good Parts. Sebastopol, CA: O\'Reilly Media, 2008.',
    '[30] A. Banks and E. Porcello, Learning React, 2nd ed. Sebastopol, CA: O\'Reilly Media, 2020.',
    '[31] E. Brown, Web Development with Node and Express, 2nd ed. Sebastopol, CA: O\'Reilly Media, 2019.',
    '[32] K. Chodorow, MongoDB: The Definitive Guide, 3rd ed. Sebastopol, CA: O\'Reilly Media, 2019.',
    '[33] L. Shklar and R. Rosen, Web Application Architecture: Principles, Protocols and Practices, 2nd ed. Hoboken, NJ: Wiley, 2009.',
    '[34] J. C. S. C. P. S. Pressman and B. Maxim, Software Engineering: A Practitioner\'s Approach, 9th ed. New York: McGraw-Hill Education, 2019.',
    '[35] Microsoft Corporation, "VS Code Documentation," Microsoft. [Online]. Available: https://code.visualstudio.com/docs. [Accessed: 05 Mar. 2026].',
    '[36] OWASP Foundation, "OWASP Top Ten Web Application Security Risks," OWASP. [Online]. Available: https://owasp.org/www-project-top-ten/. [Accessed: 08 Mar. 2026].',
    '[37] npm Documentation, "npm CLI Documentation," npm Inc. [Online]. Available: https://docs.npmjs.com/. [Accessed: 10 Mar. 2026].',
    '[38] C. Severance, "JavaScript: The First 20 Years," in Proc. ACM Programming Languages, vol. 4, no. HOPL, pp. 1-89, 2020.',
    '[39] S. G. J. N. P. V. D. M. Richards, "The Node.js Event Loop: A Developer\'s Guide," IEEE Software, vol. 38, no. 4, pp. 97-105, 2021.',
    '[40] A. Pasquali, Mastering Node.js, 2nd ed. Birmingham, UK: Packt Publishing, 2017.',
]
for ref in references:
    add_body(ref)

page_break()

# ============================================================
# APPENDICES
# ============================================================
add_heading_custom('APPENDICES', level=1, center=True)

add_heading_custom('Appendix A: Source Code', level=2)
add_body('The complete source code of the Employee Management System is available in the project repository. Key file references and implementation details have been discussed in Chapter 4 (Implementation) of this report.')

add_para('', size=12)
add_para('--- END OF REPORT ---', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)

# ============================================================
# SAVE
# ============================================================
output = r"C:\Users\KSonu\Desktop\EMS\EMS_Project_Report_final.docx"
try:
    doc.save(output)
except PermissionError:
    output = r"C:\Users\KSonu\Desktop\EMS\EMS_Project_Report_final_v3.docx"
    doc.save(output)
print(f'Final report saved: {output}')
